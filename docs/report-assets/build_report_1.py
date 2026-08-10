# -*- coding: utf-8 -*-
"""PEONIFY — University of Kigali final-year project report (Chapters 1–5).
Describes the PHP/MySQL/XAMPP implementation. Diagrams are native Word
drawings (editable shapes); the database diagram and interface figures are
authentic captures of the running system."""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import diagrams_native as dg

ASSETS = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(ASSETS, "screenshots")
DBSHOT = "/home/procell/Downloads/Screenshot (1).png"
OUT = "/home/procell/Documents/projects/luxepeony/docs/PEONIFY - FINAL PROJECT (UoK).docx"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.5
for h, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13), ("Heading 4", 12)]:
    st = doc.styles[h]
    st.font.name = "Times New Roman"; st.font.size = Pt(size)
    st.font.color.rgb = RGBColor(0, 0, 0); st.font.bold = True

FIGS, TABS = [], []

def p(text="", bold=False, italic=False, align="justify", size=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold; run.italic = italic
    if size: run.font.size = Pt(size)
    par.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                     "left": WD_ALIGN_PARAGRAPH.LEFT}.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    return par

def bullets(items):
    for it in items: doc.add_paragraph(it, style="List Bullet")

def numbered(items):
    for it in items: doc.add_paragraph(it, style="List Number")

def _cap(kind, store, title):
    n = len(store) + 1
    store.append(f"{kind} {n}: {title}")
    par = doc.add_paragraph()
    run = par.add_run(f"{kind} {n}: {title}")
    run.bold = True; run.font.size = Pt(11)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

def figcap(title): _cap("Figure", FIGS, title)
def tabcap(title): _cap("Table", TABS, title)

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def table(headers, rows, cap=None, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for r in c.paragraphs[0].runs: r.bold = True
        shade(c, "F2DCE7")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for pr in cells[i].paragraphs:
                pr.paragraph_format.space_after = Pt(2)
                for r in pr.runs: r.font.size = Pt(10.5)
    for pr in [c.paragraphs[0] for c in t.rows[0].cells]:
        for r in pr.runs: r.font.size = Pt(10.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    if cap: tabcap(cap)
    return t

def diagram(builder, cap):
    builder().add_to(doc)
    figcap(cap)

def picture(path, cap, width=6.3):
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(path, width=Inches(width))
    figcap(cap)

def screenshot(name, cap, width=6.3):
    picture(os.path.join(SHOTS, name + ".png"), cap, width)

def toc_field():
    par = doc.add_paragraph(); run = par.add_run()
    beg = OxmlElement("w:fldChar"); beg.set(qn("w:fldCharType"), "begin")
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve")
    ins.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Right-click and choose “Update Field” (or press F9) to generate the Table of Contents."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (beg, ins, sep, txt, end): run._r.append(el)

def pagebreak(): doc.add_page_break()

# ╔══════════════════════════ COVER PAGE ══════════════════════════╗
p(); p()
p("UNIVERSITY OF KIGALI (UoK)", bold=True, align="center", size=18)
p("SCHOOL OF COMPUTING AND INFORMATION TECHNOLOGY", align="center", size=13)
p("DEPARTMENT OF INFORMATION TECHNOLOGY", align="center", size=12)
p(); p()
p("PEONIFY: A WEB-BASED FLORAL E-COMMERCE AND", bold=True, align="center", size=16)
p("ORDER MANAGEMENT SYSTEM", bold=True, align="center", size=16)
p("Case study: a small floral boutique enterprise in Kigali", italic=True, align="center")
p(); p()
p("A Final Year Project Presented in Partial Fulfillment of the Requirements", align="center")
p("for the Degree of", align="center")
p("BACHELOR OF SCIENCE IN INFORMATION TECHNOLOGY", bold=True, align="center", size=13)
p(); p()
p("By", align="center")
p("[Student Full Name]", bold=True, align="center", size=13)
p("Registration Number: [Reg. No]", align="center")
p()
p("Supervisor: [Supervisor Name]", align="center")
p()
p("Kigali, Rwanda", align="center")
p("July 2026", align="center")
pagebreak()

# ╔══════════════════════════ ABSTRACT ══════════════════════════╗
doc.add_heading("ABSTRACT", level=1)
p("Peonify is a web-based floral electronic commerce and order management system developed "
  "for a small flower boutique enterprise operating in Kigali. Like the majority of "
  "micro-enterprises in the creative and artisanal sector, the boutique has historically "
  "sold its arrangements through walk-in customers and social-media direct messages. Under "
  "that way of working, orders live inside private chat conversations, payment is negotiated "
  "informally, delivery promises are made verbally, and the owner is left without durable "
  "records, without customer relationships that survive a single conversation, and without "
  "any measurable view of how the business is performing.")
p("The purpose of this project was to replace that fragile process with complete, "
  "professional commerce infrastructure that a non-technical entrepreneur can install and "
  "operate alone. The delivered system comprises a premium public storefront with searching, "
  "filtering and merchandising; an interactive bouquet builder that composes a live "
  "photographic preview of a custom arrangement; customer accounts with in-application "
  "notifications and automatic delivery reminders; a checkout that recomputes all prices on "
  "the server, records payment and issues a unique traceable order reference; a deliberately "
  "simple paid-to-delivered fulfilment workflow that both the administrator and the customer "
  "can complete; product feedback from verified customers; and an administrator dashboard "
  "presenting revenue analytics, order management, catalogue control, feedback moderation "
  "and a support inbox.")
p("The system was analysed and designed using object-oriented techniques and developed "
  "iteratively under the Agile model. It is implemented with PHP 8 and MySQL 8 on the XAMPP "
  "platform, with HTML5, CSS3 and vanilla JavaScript on the client, a stack chosen "
  "deliberately so that deployment on a Windows computer consists of copying a single folder: "
  "the database schema and demonstration catalogue create themselves on the first visit. "
  "Security follows accepted industry practice, including bcrypt password hashing, HTTP-only "
  "session cookies, anti-CSRF tokens on every form, prepared SQL statements, validated image "
  "uploads and server-side price recomputation. Functional testing covered every user "
  "journey from registration through checkout, delivery confirmation on both sides, "
  "feedback and reminders. The project demonstrates that production-shaped e-commerce "
  "infrastructure for a single merchant is achievable on an entirely open, zero-licence-cost "
  "technology stack.")
pagebreak()

# ╔══════════════════════════ DECLARATION / APPROVAL / DEDICATION ═════════╗
doc.add_heading("DECLARATION", level=1)
p("I, [Student Full Name], Registration Number [Reg. No], a student of the University of "
  "Kigali (UoK) in the School of Computing and Information Technology, do hereby declare "
  "that this final year project titled “Peonify: A Web-Based Floral E-Commerce and Order "
  "Management System” is my original work. It has never been submitted, in whole or in "
  "part, to any other university or institution of higher learning for the award of any "
  "degree or diploma. All materials and sources of information used in this work have been "
  "duly acknowledged and referenced.")
p(); p("Signature: ………………………………                Date: ………………………………")
pagebreak()

doc.add_heading("APPROVAL", level=1)
p("This is to certify that the final year project titled “Peonify: A Web-Based Floral "
  "E-Commerce and Order Management System” was carried out by [Student Full Name], "
  "Registration Number [Reg. No], under my supervision and guidance, and that it is now "
  "ready for submission and examination with my approval.")
p(); p("Supervisor: [Supervisor Name]")
p("Signature: ………………………………                Date: ………………………………")
pagebreak()

doc.add_heading("DEDICATION", level=1)
p("This work is dedicated to my beloved family, whose patience, sacrifices and constant "
  "encouragement have carried me through every stage of my studies, and to my friends and "
  "classmates who stood by me during the long days and nights this project required.")
pagebreak()

doc.add_heading("ACKNOWLEDGEMENTS", level=1)
p("First and foremost, I thank the Almighty God for the gift of life, health and strength "
  "throughout the period of my studies and of this project.")
p("I express my sincere gratitude to the University of Kigali, and in particular to the "
  "School of Computing and Information Technology, for providing the academic environment "
  "and resources within which this work was possible. I am especially indebted to my "
  "supervisor, [Supervisor Name], whose guidance, patience and constructive criticism "
  "shaped this project from an idea into a working system and a documented study.")
p("I also thank the flower entrepreneurs who generously shared their daily working "
  "experience during the requirements-gathering phase; their openness about the practical "
  "difficulties of selling flowers through informal channels grounded this project in a "
  "real problem. Finally, I thank my family and friends for their unconditional support.")
pagebreak()

# ╔══════════════════════════ TOC & LISTS ══════════════════════════╗
doc.add_heading("TABLE OF CONTENTS", level=1)
toc_field()
pagebreak()
LOF_HEAD = doc.add_heading("LIST OF FIGURES", level=1)
LOF_PAR = p("(generated below)")
doc.add_heading("LIST OF TABLES", level=1)
LOT_PAR = p("(generated below)")
pagebreak()
doc.add_heading("LIST OF ABBREVIATIONS AND ACRONYMS", level=1)
for ab in ["AJAX – Asynchronous JavaScript And XML",
           "API – Application Programming Interface",
           "CDN – Content Delivery Network",
           "CRUD – Create, Read, Update, Delete",
           "CSRF – Cross-Site Request Forgery",
           "CSS – Cascading Style Sheets",
           "ERD – Entity Relationship Diagram",
           "FR – Functional Requirement",
           "HTML – HyperText Markup Language",
           "HTTP – HyperText Transfer Protocol",
           "ICT – Information and Communication Technology",
           "IT – Information Technology",
           "JS – JavaScript",
           "JSON – JavaScript Object Notation",
           "NFR – Non-Functional Requirement",
           "OOAD – Object-Oriented Analysis and Design",
           "PDO – PHP Data Objects",
           "PHP – PHP: Hypertext Preprocessor",
           "SQL – Structured Query Language",
           "UAT – User Acceptance Testing",
           "UI – User Interface",
           "UML – Unified Modeling Language",
           "UoK – University of Kigali",
           "XAMPP – Cross-platform Apache MySQL PHP Perl"]:
    doc.add_paragraph(ab, style="List Bullet")
pagebreak()

# ╔══════════════════════════ CHAPTER 1 ══════════════════════════╗
doc.add_heading("CHAPTER 1", level=1)
doc.add_heading("GENERAL INTRODUCTION", level=1)

doc.add_heading("1.1 Introduction", level=2)
p("In today’s digital economy, businesses of every size are expected to serve their "
  "customers online. Electronic commerce has evolved from a competitive advantage into a "
  "baseline expectation: customers assume that they can browse a catalogue at any hour, "
  "compare prices transparently, pay securely, choose when a purchase will reach them, and "
  "follow the progress of their order without making a phone call. Enterprises that cannot "
  "offer this experience increasingly lose customers to those that can, regardless of the "
  "quality of the underlying product.")
p("This expectation now reaches even the smallest businesses. A neighbourhood florist "
  "competes, in the mind of the customer, with the polished ordering experience of "
  "international platforms. Yet the technical and financial barriers that once made such "
  "an experience unattainable for a micro-enterprise have largely disappeared: open-source "
  "server software, free relational databases and modern browsers make it possible to "
  "deliver professional commerce infrastructure at zero licence cost. This project applies "
  "that opportunity to a concrete case: it presents the analysis, design, implementation "
  "and testing of Peonify, a web-based floral e-commerce and order management system "
  "developed for a small flower boutique, using PHP, MySQL, HTML, CSS and vanilla "
  "JavaScript on the XAMPP platform.")

doc.add_heading("1.2 Background of the Study", level=2)
p("Flowers occupy a distinctive position in retail commerce. They are purchased primarily "
  "for emotional occasions — anniversaries, weddings, graduations, apologies, condolences "
  "— which makes reliability of delivery at least as important as the product itself. They "
  "are also highly perishable: an arrangement assembled in the morning must reach its "
  "recipient within hours to arrive at its best. Together these characteristics make "
  "timing and trust the two currencies of the floral trade.")
p("Globally, floral retail has moved decisively online. The flower delivery service market "
  "was estimated at USD 7.6 billion in 2024 and is projected to reach USD 11.27 billion by "
  "2030, growing at approximately seven percent per year (Grand View Research, 2024). "
  "Industry analyses report that online channels already account for roughly one third of "
  "all flower sales and that this share continues to grow annually (BusinessDojo, 2025). "
  "Established platforms such as 1-800-Flowers in the United States, Bloom & Wild in "
  "Europe and Floward in the Middle East have demonstrated that customers will confidently "
  "order premium, perishable arrangements from a screen when the experience communicates "
  "trust: authentic photography, transparent prices, precise delivery windows and visible "
  "order progress.")
p("In Rwanda, the environment for such systems is increasingly favourable. The country has "
  "made sustained investments in digital infrastructure and digital literacy, and a "
  "growing share of the population engages with online services primarily through mobile "
  "devices. Small businesses have responded by adopting the tools most immediately "
  "available to them: social-media platforms. Micro-entrepreneurs worldwide use Instagram, "
  "WhatsApp and similar applications as virtual storefronts because they are free and "
  "familiar (Mastercard Strive, 2024). The floral boutique studied in this project is "
  "typical of this pattern: talented at its craft, visible on social media, but entirely "
  "dependent on informal, manual processes for everything that happens after a customer "
  "says “I would like to order.”")
p("Studies of chat-based commerce consistently document the consequences of this "
  "dependency: fragmented customer journeys, manual order management, unstructured "
  "payments and the absence of any delivery discipline (PwC, 2024). Consumer trust is a "
  "further constraint — nearly half of surveyed consumers worry that purchases made "
  "through social media will not be protected or refunded if something goes wrong (Tidio, "
  "2024). The gap between what small florists can offer and what customers have learned "
  "to expect is precisely the gap this project addresses.")

doc.add_heading("1.3 Problem Statement", level=2)
p("The floral boutique that forms the case study of this project currently manages its "
  "entire commercial operation through informal channels. Product photographs are posted "
  "to social media; interested buyers send private messages; prices, delivery locations "
  "and delivery times are negotiated conversation by conversation; payment is collected in "
  "cash on delivery or through person-to-person mobile transfers; and the fulfilment of "
  "each order depends on the owner remembering it. This way of working is time-consuming "
  "and error-prone, and it produces no permanent record of what was sold, to whom, for "
  "how much, or whether it was ever delivered.")
p("As the volume of orders grows, the weaknesses of the manual process compound. Orders "
  "are forgotten or confused with one another; disputed prices cannot be resolved because "
  "no authoritative record exists; regular customers must reintroduce themselves with "
  "every purchase; and the owner cannot answer even elementary business questions such as "
  "how much revenue the boutique earned last month or which arrangements sell best. In "
  "the absence of a system, the ceiling of the business is the memory and availability of "
  "one person.")
doc.add_heading("1.3.1 Specific Problems", level=3)
bullets([
    "Ephemeral visibility: social-media posts disappear from followers’ feeds within "
    "hours, so the catalogue must be re-advertised continuously and has no permanent, "
    "searchable home.",
    "Informal ordering: orders captured in private chat threads suffer wrong addresses, "
    "missed dates and disputed prices, and can be neither tracked nor audited.",
    "Unstructured payments: cash and person-to-person transfers offer no receipts, no "
    "reconciliation against orders and no protection for either party.",
    "Absent customer relationships: without accounts, the business cannot notify "
    "customers of order progress, remind them of upcoming deliveries, or accumulate the "
    "verified feedback that builds trust for future buyers.",
    "No business intelligence: the owner has no view of revenue over time, of the order "
    "pipeline, or of product performance on which to base stocking and pricing decisions.",
])

doc.add_heading("1.4 Motivation of the Study", level=2)
p("The decision to develop Peonify is driven by the conviction that the benefits of "
  "digital commerce should not be reserved for large enterprises. The specific motivations "
  "of the study are threefold.")
p("To the University of Kigali, this study aligns with the university’s mission of "
  "producing graduates who apply information technology to solve practical problems in "
  "the surrounding community. It demonstrates the complete software engineering lifecycle "
  "— analysis, design, implementation, testing and documentation — applied to a genuine "
  "small-business problem rather than an abstract exercise.")
p("To the flower entrepreneur, the study delivers a working, zero-licence-cost system "
  "that professionalises daily operations: a permanent catalogue, verifiable payment "
  "records, disciplined deliveries, lasting customer relationships and actionable "
  "analytics — all operable without technical skills and installable by copying a single "
  "folder onto a Windows computer running XAMPP.")
p("To the student researcher, the project offered the opportunity to confront the "
  "realities of production software: security hardening, data integrity, usability for "
  "non-technical users, and the discipline of testing every user journey rather than only "
  "the successful path.")

doc.add_heading("1.5 Objectives of the Study", level=2)
doc.add_heading("1.5.1 General Objective", level=3)
p("The general objective of this study is to design and develop a web-based floral "
  "e-commerce and order management system that enables a flower entrepreneur to sell "
  "premium arrangements online with recorded payments, scheduled deliveries and "
  "transparent order progress, while remaining simple enough for a non-technical owner "
  "to install and operate end to end.")
doc.add_heading("1.5.2 Specific Objectives", level=3)
numbered([
    "To analyse the current, manual, social-media-based selling process of a small floral "
    "business and to model its shortcomings.",
    "To design a premium, responsive storefront — including an interactive bouquet "
    "builder with a live photographic preview — that converts visitors into buyers.",
    "To implement secure account management with distinct customer and administrator "
    "roles, using bcrypt password hashing, PHP sessions and anti-CSRF protection.",
    "To implement a checkout and payment-recording module that recomputes order totals on "
    "the server, marks orders as paid and issues unique, traceable order references.",
    "To implement a deliberately simple paid-to-delivered fulfilment workflow with "
    "in-application notifications and automatic delivery reminders for both parties.",
    "To provide the administrator with analytics — a thirty-day revenue chart, the order "
    "pipeline and top-selling products — together with complete catalogue control over "
    "products, categories, collections and promotional discounts.",
    "To test the developed system functionally across every user journey and document "
    "the results.",
])

doc.add_heading("1.6 Scope of the Study", level=2)
doc.add_heading("1.6.1 Content Scope", level=3)
p("The system covers the complete commercial cycle of a single floral boutique: the "
  "public storefront and merchandising; the bouquet builder; customer registration, "
  "authentication and profiles; the cart and checkout with delivery scheduling; payment "
  "recording with unique order references; order fulfilment and confirmation; "
  "notifications and delivery reminders; product feedback and its moderation; customer "
  "support messaging; and the administrator dashboard with analytics and catalogue "
  "management.")
doc.add_heading("1.6.2 Geographical Scope", level=3)
p("The case study targets a floral boutique operating in Kigali, Rwanda. Kigali was "
  "chosen because it concentrates the customer base most likely to order flowers online, "
  "enjoys the country’s highest levels of internet penetration and digital literacy, and "
  "hosts the events — weddings, conferences, graduations — that drive premium floral "
  "demand.")
doc.add_heading("1.6.3 Technological Scope", level=3)
p("The system is implemented for the XAMPP platform: Apache as the web server, PHP 8 as "
  "the server-side language, MySQL 8 as the relational database, and HTML5, CSS3 and "
  "vanilla JavaScript on the client. No JavaScript frameworks and no build tooling are "
  "used, so the application runs directly from the web server’s document root.")
doc.add_heading("1.6.4 Out of Scope", level=3)
bullets([
    "Multi-vendor marketplace functionality — the system serves a single merchant by design.",
    "Native mobile applications — the responsive web interface serves mobile browsers.",
    "Courier fleet management and third-party logistics integration.",
    "Subscription or recurring billing.",
    "Integration with an external online payment gateway, which is identified as future work.",
])

doc.add_heading("1.7 Methods and Techniques Used in the Study", level=2)
doc.add_heading("1.7.1 Data Collection Techniques", level=3)
p("Three complementary techniques were used to establish the requirements of the system. "
  "Documentation review examined how established floral platforms — 1-800-Flowers, "
  "Floward and Bloom & Wild — and general marketplaces present products, schedule "
  "deliveries and communicate order progress, in order to identify the interaction "
  "patterns customers already understand. Direct observation followed the day-to-day "
  "selling activity of small florists on Instagram and WhatsApp, recording how orders are "
  "initiated, negotiated and too often lost. Finally, informal interviews with the target "
  "entrepreneur established the practical constraints of the business: no technical "
  "staff, a single computer running Windows, and no budget for software licences or "
  "monthly platform fees.")
doc.add_heading("1.7.2 System Development Methodology", level=3)
p("Development followed the Agile iterative model. The system was built in vertical "
  "slices — database, server logic, user interface and test for one feature at a time — "
  "and each slice was demonstrated and refined before the next began. This approach "
  "proved essential: several design decisions, most notably the simplification of the "
  "fulfilment pipeline into a single paid-to-delivered transition and the addition of "
  "customer delivery confirmation, emerged directly from iteration feedback rather than "
  "from the initial design.")
doc.add_heading("1.7.3 Analysis and Design Techniques", level=3)
p("Object-oriented analysis and design (OOAD) was used to model the system. Use case "
  "modelling captured functional requirements from the perspective of each actor; class "
  "modelling identified the principal entities and their responsibilities; sequence "
  "modelling documented the checkout interaction; and relational design mapped the model "
  "onto a normalised MySQL schema. The Unified Modeling Language (UML) notation is used "
  "throughout Chapter 3.")

doc.add_heading("1.8 Significance of the Study", level=2)
p("For the boutique, the significance is immediate and operational: a permanent, "
  "searchable catalogue replaces ephemeral posts; recorded payments with unique "
  "references replace unverifiable transfers; scheduled, reminded deliveries replace "
  "promises made from memory; and a dashboard finally answers the questions every owner "
  "asks — what is selling, what is pending, and how is revenue moving.")
p("For the wider small-business community, the project provides a replicable template. "
  "Because the entire stack is open-source and the deployment procedure is a folder copy, "
  "the same architecture can serve bakeries, tailors, artisans and any other single-"
  "merchant business that today depends on chat-based selling.")
p("For the academic community, the project documents a complete, security-conscious "
  "implementation of a real commerce system on the classic PHP/MySQL stack that dominates "
  "small-business hosting worldwide, demonstrating that rigorous engineering practice is "
  "compatible with the simplest and most accessible tools.")

doc.add_heading("1.9 Expected Results", level=2)
bullets([
    "A permanent, professionally merchandised online boutique replacing ephemeral "
    "social-media posts.",
    "Recorded, reconcilable payments with unique order references and automatic "
    "notifications to both parties.",
    "A complete and auditable order history, with delivery reminders sent twenty-four "
    "hours and five hours before each delivery window.",
    "Verified customer feedback displayed on products and on the landing page, with "
    "administrative moderation.",
    "Actionable analytics for the entrepreneur: daily revenue over thirty days, the "
    "order pipeline and top-selling arrangements.",
    "A deployment procedure achievable by a non-developer: install XAMPP, copy one "
    "folder, open the browser.",
])

doc.add_heading("1.10 Definition of Key Terms", level=2)
bullets([
    "E-commerce: the buying and selling of goods or services over the internet, "
    "including the payment and fulfilment processes that complete each transaction.",
    "Order management: the recording, tracking and completion of customer orders from "
    "placement through delivery, with an auditable history of every state change.",
    "Storefront: the public, customer-facing portion of an e-commerce system where "
    "products are presented, discovered and added to a cart.",
    "Checkout: the process in which a customer provides delivery details, the system "
    "computes the authoritative total, payment is recorded and an order is created.",
    "Fulfilment: the operational completion of a paid order, ending with delivery to "
    "the customer and confirmation within the system.",
    "Session: the server-maintained association between a signed-in user and their "
    "subsequent requests, carried by a protected browser cookie.",
    "Seeding: the automatic insertion of initial data — accounts, taxonomy and "
    "demonstration products — when a system is installed.",
])

doc.add_heading("1.11 Organization of the Report", level=2)
p("The remainder of this report is organised as follows. Chapter two analyses the "
  "existing manual system, models its process, articulates its problems and derives the "
  "functional and non-functional requirements of the proposed solution. Chapter three "
  "presents the requirements analysis and design of the new system using UML models, the "
  "database schema and the system architecture. Chapter four discusses the technical "
  "implementation — the technologies used and their justification, the realisation of "
  "each module, the security measures applied — and presents the developed system through "
  "authentic screen captures together with the testing performed. Chapter five concludes "
  "the study and offers recommendations and directions for future work.")
pagebreak()

# ╔══════════════════════════ CHAPTER 2 ══════════════════════════╗
doc.add_heading("CHAPTER 2", level=1)
doc.add_heading("ANALYSIS OF THE EXISTING SYSTEM", level=1)

doc.add_heading("2.1 Introduction", level=2)
p("Before a new system can be designed responsibly, the system it replaces must be "
  "understood in detail — not only its mechanics but the reasons it persists and the "
  "precise points at which it fails. This chapter describes the environment and history "
  "of the boutique’s current way of working, models the existing manual process, "
  "analyses its problems, presents the proposed solution, and derives the functional and "
  "non-functional requirements that guided the design in Chapter 3.")

doc.add_heading("2.2 Description of the Existing System Environment", level=2)
doc.add_heading("2.2.1 Historical Background", level=3)
p("The boutique began as a home-based flower business promoted through the owner’s "
  "personal social-media accounts. Early orders came from friends and acquaintances, for "
  "whom informality was natural: a message, an agreed price, a delivery arranged in "
  "person. As word spread, the owner adopted a business Instagram profile and WhatsApp "
  "Business for order taking, and order volume grew beyond the circle of personal "
  "acquaintances. The tools, however, did not change. All records — orders, payments, "
  "delivery promises — remained where they were born: inside chat histories and, "
  "occasionally, a personal notebook. No computerised system of any kind has ever been "
  "used in the business.")
doc.add_heading("2.2.2 Current Operating Context", level=3)
p("Today the boutique sources fresh stems from local growers, assembles arrangements in "
  "a small workshop, and sells through two channels: walk-in customers, and social-media "
  "direct messages that now account for the majority of orders. Deliveries within Kigali "
  "are made personally by the owner or by informal courier arrangements agreed order by "
  "order. The owner operates alone, with occasional family help during peak periods such "
  "as Valentine’s season and graduation weeks — precisely the periods in which the "
  "manual process is most likely to fail.")

doc.add_heading("2.3 Description of the Existing System", level=2)
p("The existing “system” is a sequence of manual activities threaded through chat "
  "applications. A typical order proceeds as follows. The owner photographs finished "
  "arrangements and posts them to Instagram with an invitation to order by direct "
  "message. An interested buyer writes privately, often beginning a negotiation over "
  "price, stem selection, wrapping, delivery place and time. When agreement is reached — "
  "sometimes across dozens of messages spread over hours — the buyer either promises "
  "cash on delivery or sends a mobile-money transfer, of which the only record is the "
  "transfer message itself. The owner then writes the delivery into memory or a "
  "notebook, prepares the arrangement on the agreed morning, and delivers it. If the "
  "buyer later wishes to order again, the entire conversation begins from zero; nothing "
  "about them — address, preferences, history — has been retained anywhere searchable.")
p("Feedback, when it exists, arrives as a private thank-you message seen only by the "
  "owner. It builds no public trust and influences no future buyer. Equally, complaints "
  "arrive privately and leave no trace that could improve the process.")

doc.add_heading("2.4 Analysis of the Existing System", level=2)
doc.add_heading("2.4.1 Modeling of the Existing System", level=3)
p("Figure 1 models the existing process as a flow of five manual activities. Beneath "
  "each activity, the analysis notes the characteristic failure it introduces. The model "
  "makes visible what the participants experience daily: every stage depends on human "
  "memory and goodwill, and no stage produces a durable, verifiable record.")
diagram(dg.existing_process, "Model of the existing manual selling process")
p("Reading the model from left to right: visibility exists only while a post is recent; "
  "the order enters a private thread indistinguishable from dozens of others; the "
  "commercial terms are improvised rather than standing; the payment leaves no "
  "reconcilable record; and the delivery depends on recollection. A failure at any stage "
  "is invisible until a customer complains.")

doc.add_heading("2.5 Problems of the Existing System", level=2)
p("The analysis identifies six systemic problems, each traceable to a stage of the model "
  "above.")
bullets([
    "No permanent catalogue: every product must be re-posted to remain visible, and a "
    "buyer who saw an arrangement last week has no way to find it today.",
    "Scattered, unauditable orders: commitments live in chat threads where they are "
    "easily forgotten, confused between customers, or lost when a device changes.",
    "Unverifiable payments: cash and person-to-person transfers cannot be reconciled "
    "against orders; disputes reduce to one person’s word against another’s.",
    "Undisciplined delivery: with no scheduling and no reminders, deliveries are missed "
    "or late precisely during high-demand periods, damaging the trust on which an "
    "emotional purchase depends.",
    "No accumulated customer relationships: addresses, preferences and history evaporate "
    "with each conversation, making loyalty invisible and repeat purchase needlessly "
    "difficult.",
    "No reporting of any kind: the owner cannot state monthly revenue, count pending "
    "orders, or identify the best-selling arrangement other than by impression.",
])

doc.add_heading("2.6 Review of Similar Systems", level=2)
p("Before designing a replacement, existing ways of selling flowers online were "
  "reviewed to identify both the interaction patterns customers already understand and "
  "the gaps that justify a purpose-built system for a single boutique.")
doc.add_heading("2.6.1 Social-Media Selling (Instagram and WhatsApp Business)", level=3)
p("Social platforms provide instant reach at no cost and remain the de facto standard "
  "for small florists — indeed, they are the existing system of this study. Their "
  "strengths end at visibility: they provide no catalogue structure, no ordering, no "
  "payment records and no delivery discipline, and the platform’s algorithm rather "
  "than the merchant decides who sees each product. They demonstrate the audience but "
  "not the infrastructure.")
doc.add_heading("2.6.2 General Marketplaces (Alibaba, Kikuu, Jumia)", level=3)
p("General marketplaces contribute the commerce mechanics customers already know: "
  "structured listings, carts, discount badges, best-seller sections and buyer "
  "reviews. Peonify deliberately borrows these patterns. As a home for a premium "
  "perishable product, however, marketplaces fall short: a bouquet requiring same-day "
  "delivery within a chosen time window does not fit commodity logistics, the product "
  "drowns among unrelated goods, and the merchant surrenders both brand identity and "
  "the customer relationship to the platform.")
doc.add_heading("2.6.3 Hosted Store Builders (Shopify and Similar)", level=3)
p("Hosted builders give a merchant a branded store with themes and plugins, and prove "
  "that non-developers can run online shops. Their limits for this case are economic "
  "and functional: recurring monthly fees are significant for a micro-enterprise; "
  "florist-specific needs — delivery windows, bouquet configuration, delivery "
  "reminders — require additional paid applications; and the administration surface "
  "remains generic and complex relative to the owner’s actual daily tasks.")
doc.add_heading("2.6.4 Dedicated Floral Platforms (1-800-Flowers, Floward, Bloom & Wild)", level=3)
p("The dedicated platforms represent the customer-experience benchmark: authentic "
  "photography, curated collections, precise delivery scheduling and visible order "
  "progress. Floward, for example, built its regional leadership on reliable same-day "
  "delivery within tight windows. These platforms are proprietary businesses rather "
  "than software an independent florist can adopt; joining them, where possible at "
  "all, means surrendering margin and brand. Peonify adopts their proven experience "
  "patterns in software the entrepreneur owns outright.")
doc.add_heading("2.6.5 Comparative Summary", level=3)
table(["System", "Strengths", "Weaknesses for this case"],
    [["Social-media selling", "Free; instant reach; visual medium suits flowers", "No orders, payments, scheduling or records; algorithm-dependent visibility"],
     ["General marketplaces", "Familiar commerce mechanics; integrated payments; reviews", "Commodity logistics unfit for perishables; brand and customer data surrendered"],
     ["Hosted store builders", "Branded store without coding; plugin ecosystem", "Recurring fees; florist features cost extra; complex generic administration"],
     ["Dedicated floral platforms", "Benchmark delivery experience and trust", "Closed businesses, not adoptable software; margin and brand loss"],
     ["Peonify (proposed)", "Purpose-built florist workflows; owner-operated; zero licence cost; self-installing", "Single merchant by design; placeholder imagery until owner photography is added"]],
    cap="Comparison of existing approaches with the proposed system", widths=[1.5, 2.4, 2.5])

doc.add_heading("2.7 Proposed Solution", level=2)
p("The proposed solution is Peonify, a self-contained web-based e-commerce and order "
  "management system installed on the XAMPP platform. Peonify addresses each identified "
  "problem directly. A permanent, searchable storefront with professional merchandising "
  "replaces ephemeral posts. Account-based ordering through a structured checkout — with "
  "recipient details, a delivery date and a selectable time window — replaces chat "
  "negotiation. Payment is recorded at checkout with the order total recomputed on the "
  "server and a unique order reference issued for every purchase, replacing unverifiable "
  "transfers. A deliberately simple fulfilment workflow, automatic notifications on both "
  "sides and delivery reminders twenty-four hours and five hours before the chosen "
  "window replace memory-based delivery. Customer accounts retain addresses and order "
  "history, and verified customers may rate and review products, accumulating public "
  "trust. Finally, an administrator dashboard presents revenue analytics, the order "
  "pipeline and catalogue management in plain language designed for a non-technical "
  "owner.")
p("Critically for this case, the system is engineered for autonomous deployment: on its "
  "first visit it creates its own database, all tables and a demonstration catalogue, so "
  "that installation on a Windows computer consists of installing XAMPP and copying one "
  "folder into the web server’s document root.")

doc.add_heading("2.8 System Requirements", level=2)
p("The requirements below were derived from the problem analysis and validated against "
  "the entrepreneur’s working needs. They are stated as testable capabilities; the test "
  "results appear in Chapter 4.")
doc.add_heading("2.8.1 Functional Requirements", level=3)
table(["Code", "Requirement", "Description"],
    [["FR1", "User registration and authentication", "The system shall allow users to register and sign in with an email address and password. Passwords shall be stored only as bcrypt hashes; sessions shall use HTTP-only cookies; the administrator account shall be seeded automatically on first run."],
     ["FR2", "Role-based access control", "The system shall distinguish customers from the administrator. Administrators shall not be able to purchase; customers shall not be able to reach management functions; each role shall be directed to its own dashboard."],
     ["FR3", "Product catalogue management", "The administrator shall create, edit and delete products — with photo upload, price, category, collection, stock status and an optional promotional discount — through a modal form, with changes visible in the storefront immediately."],
     ["FR4", "Storefront and discovery", "Customers shall browse with text search, category chips, a collection filter, a price-range filter, a sale-only filter, sorting and pagination. The landing page shall merchandise best sellers, new arrivals, flash sales, collections and customer testimonials."],
     ["FR5", "Bouquet builder", "The system shall offer a four-step configurator (size, focal flower, foliage, packaging) with a live photographic preview and a per-step price breakdown; the chosen configuration shall be stored with the order line."],
     ["FR6", "Cart and checkout", "The cart shall persist in the browser. Checkout shall require an account, pre-fill recipient details from the profile, and capture the delivery address, date, time window and an optional gift note."],
     ["FR7", "Payment recording", "At checkout the system shall recompute every price on the server from the database, record the payment, mark the order as paid, and issue a unique human-readable reference (e.g., PNY-A1B2C3)."],
     ["FR8", "Fulfilment workflow", "Orders shall move from paid to delivered through a single action: the administrator pressing Deliver (with a confirmation dialog) or the receiving customer pressing “I received it”. Every transition shall be logged as a timestamped event."],
     ["FR9", "Notifications and reminders", "Both roles shall receive in-application notifications with unread badges and per-item and mark-all read controls. Customers shall be reminded twenty-four hours and five hours before their delivery window."],
     ["FR10", "Feedback and support", "Verified customers shall rate and review products (one review per customer per product, editable). The public contact form and the account support tab shall deliver messages to the administrator’s inbox. The administrator shall be able to remove inappropriate reviews."],
     ["FR11", "Analytics dashboard", "The administrator dashboard shall chart daily revenue over the last thirty days and the order pipeline, list top-selling products, and display headline statistics with compact money formatting."],
     ["FR12", "Profile management", "Both roles shall manage a profile: a click-to-upload photograph, name, phone, and a saved delivery address and city used to pre-fill checkout; password change shall require the current password and a matching confirmation."]],
    cap="Functional requirements of the proposed system", widths=[0.55, 1.55, 4.3])
doc.add_heading("2.8.2 Non-Functional Requirements", level=3)
table(["Code", "Requirement", "Description"],
    [["NFR1", "Usability", "All interfaces shall use plain language; the administrator dashboard shall be operable by a non-developer; destructive actions shall require confirmation."],
     ["NFR2", "Security", "The system shall apply bcrypt hashing (cost 12), HTTP-only SameSite session cookies, anti-CSRF tokens on every form, prepared SQL statements exclusively, image-only validated uploads limited to five megabytes, and server-side price recomputation."],
     ["NFR3", "Performance", "Long lists shall be paginated; images shall load lazily; each request shall use a single pooled database connection; interface libraries shall be served from CDNs."],
     ["NFR4", "Portability", "The system shall run on a default XAMPP installation on Windows without configuration; the database shall create and seed itself on first run."],
     ["NFR5", "Reliability", "Order creation shall be transactional; totals shall always be recomputed on the server before an order is stored."],
     ["NFR6", "Maintainability", "The code shall be organised as one page-controller per screen with shared include files; schema migration and seeding shall be idempotent."],
     ["NFR7", "Responsiveness", "All screens, including the administrator dashboard, shall adapt to mobile viewports."],
     ["NFR8", "Availability", "The system shall operate wholly on the local server without external service dependencies for its core flows."]],
    cap="Non-functional requirements of the proposed system", widths=[0.6, 1.3, 4.5])

doc.add_heading("2.9 Feasibility of the Proposed Solution", level=2)
p("Technical feasibility: every component of the chosen stack — Apache, PHP 8, MySQL 8, "
  "HTML5, CSS3 and vanilla JavaScript — is mature, freely available and documented "
  "exhaustively; no experimental technology is required.")
p("Operational feasibility: the administrator interface was designed with, and validated "
  "against, a non-technical user. Every operational task — adding a product, delivering "
  "an order, answering a message — is a single screen with plain-language labels and "
  "confirmation dialogs.")
p("Economic feasibility: the entire software stack carries zero licence cost. The only "
  "recurring costs of a production deployment are a domain name and basic hosting, both "
  "of which are widely affordable to a functioning boutique.")
pagebreak()
