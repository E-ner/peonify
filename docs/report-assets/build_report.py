# -*- coding: utf-8 -*-
"""PEONIFY — University of Kigali final-year project report (Chapters 1–5).
Describes the PHP/MySQL/XAMPP implementation. Diagrams are native Word
drawings (editable shapes); the database diagram and interface figures are
authentic captures of the running system."""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import diagrams_native as dg

ASSETS = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(ASSETS, "screenshots")
DBSHOT = "/home/procell/Downloads/Screenshot (1).png"
OUT = "/home/procell/Documents/projects/luxepeony/docs/PEONIFY - FINAL PROJECT (UoK).docx"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.5
for h, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13), ("Heading 4", 12)]:
    st = doc.styles[h]
    st.font.name = "Times New Roman"; st.font.size = Pt(size)
    st.font.color.rgb = RGBColor(0, 0, 0); st.font.bold = True

FIGS, TABS = [], []

def p(text="", bold=False, italic=False, align="justify", size=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold; run.italic = italic
    if size: run.font.size = Pt(size)
    par.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                     "left": WD_ALIGN_PARAGRAPH.LEFT}.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    return par

def bullets(items):
    for it in items: doc.add_paragraph(it, style="List Bullet")

def numbered(items):
    for it in items: doc.add_paragraph(it, style="List Number")

def _cap(kind, store, title):
    n = len(store) + 1
    store.append(f"{kind} {n}: {title}")
    par = doc.add_paragraph()
    run = par.add_run(f"{kind} {n}: {title}")
    run.bold = True; run.font.size = Pt(11)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

def figcap(title): _cap("Figure", FIGS, title)
def tabcap(title): _cap("Table", TABS, title)

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def table(headers, rows, cap=None, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for r in c.paragraphs[0].runs: r.bold = True
        shade(c, "F2DCE7")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for pr in cells[i].paragraphs:
                pr.paragraph_format.space_after = Pt(2)
                for r in pr.runs: r.font.size = Pt(10.5)
    for pr in [c.paragraphs[0] for c in t.rows[0].cells]:
        for r in pr.runs: r.font.size = Pt(10.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    if cap: tabcap(cap)
    return t

def diagram(builder, cap):
    builder().add_to(doc)
    figcap(cap)

def picture(path, cap, width=6.3):
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(path, width=Inches(width))
    figcap(cap)

def screenshot(name, cap, width=6.3):
    picture(os.path.join(SHOTS, name + ".png"), cap, width)

def toc_field():
    par = doc.add_paragraph(); run = par.add_run()
    beg = OxmlElement("w:fldChar"); beg.set(qn("w:fldCharType"), "begin")
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve")
    ins.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Right-click and choose “Update Field” (or press F9) to generate the Table of Contents."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (beg, ins, sep, txt, end): run._r.append(el)

def pagebreak(): doc.add_page_break()

# ╔══════════════════════════ COVER PAGE ══════════════════════════╗
p(); p()
p("UNIVERSITY OF KIGALI (UoK)", bold=True, align="center", size=18)
p("SCHOOL OF COMPUTING AND INFORMATION TECHNOLOGY", align="center", size=13)
p("DEPARTMENT OF INFORMATION TECHNOLOGY", align="center", size=12)
p(); p()
p("PEONIFY: A WEB-BASED FLORAL E-COMMERCE AND", bold=True, align="center", size=16)
p("ORDER MANAGEMENT SYSTEM", bold=True, align="center", size=16)
p("Case study: a small floral boutique enterprise in Kigali", italic=True, align="center")
p(); p()
p("A Final Year Project Presented in Partial Fulfillment of the Requirements", align="center")
p("for the Degree of", align="center")
p("BACHELOR OF SCIENCE IN INFORMATION TECHNOLOGY", bold=True, align="center", size=13)
p(); p()
p("By", align="center")
p("[Student Full Name]", bold=True, align="center", size=13)
p("Registration Number: [Reg. No]", align="center")
p()
p("Supervisor: [Supervisor Name]", align="center")
p()
p("Kigali, Rwanda", align="center")
p("July 2026", align="center")
pagebreak()

# ╔══════════════════════════ ABSTRACT ══════════════════════════╗
doc.add_heading("ABSTRACT", level=1)
p("Peonify is a web-based floral electronic commerce and order management system developed "
  "for a small flower boutique enterprise operating in Kigali. Like the majority of "
  "micro-enterprises in the creative and artisanal sector, the boutique has historically "
  "sold its arrangements through walk-in customers and social-media direct messages. Under "
  "that way of working, orders live inside private chat conversations, payment is negotiated "
  "informally, delivery promises are made verbally, and the owner is left without durable "
  "records, without customer relationships that survive a single conversation, and without "
  "any measurable view of how the business is performing.")
p("The purpose of this project was to replace that fragile process with complete, "
  "professional commerce infrastructure that a non-technical entrepreneur can install and "
  "operate alone. The delivered system comprises a premium public storefront with searching, "
  "filtering and merchandising; an interactive bouquet builder that composes a live "
  "photographic preview of a custom arrangement; customer accounts with in-application "
  "notifications and automatic delivery reminders; a checkout that recomputes all prices on "
  "the server, records payment and issues a unique traceable order reference; a deliberately "
  "simple paid-to-delivered fulfilment workflow that both the administrator and the customer "
  "can complete; product feedback from verified customers; and an administrator dashboard "
  "presenting revenue analytics, order management, catalogue control, feedback moderation "
  "and a support inbox.")
p("The system was analysed and designed using object-oriented techniques and developed "
  "iteratively under the Agile model. It is implemented with PHP 8 and MySQL 8 on the XAMPP "
  "platform, with HTML5, CSS3 and vanilla JavaScript on the client, a stack chosen "
  "deliberately so that deployment on a Windows computer consists of copying a single folder: "
  "the database schema and demonstration catalogue create themselves on the first visit. "
  "Security follows accepted industry practice, including bcrypt password hashing, HTTP-only "
  "session cookies, anti-CSRF tokens on every form, prepared SQL statements, validated image "
  "uploads and server-side price recomputation. Functional testing covered every user "
  "journey from registration through checkout, delivery confirmation on both sides, "
  "feedback and reminders. The project demonstrates that production-shaped e-commerce "
  "infrastructure for a single merchant is achievable on an entirely open, zero-licence-cost "
  "technology stack.")
pagebreak()

# ╔══════════════════════════ DECLARATION / APPROVAL / DEDICATION ═════════╗
doc.add_heading("DECLARATION", level=1)
p("I, [Student Full Name], Registration Number [Reg. No], a student of the University of "
  "Kigali (UoK) in the School of Computing and Information Technology, do hereby declare "
  "that this final year project titled “Peonify: A Web-Based Floral E-Commerce and Order "
  "Management System” is my original work. It has never been submitted, in whole or in "
  "part, to any other university or institution of higher learning for the award of any "
  "degree or diploma. All materials and sources of information used in this work have been "
  "duly acknowledged and referenced.")
p(); p("Signature: ………………………………                Date: ………………………………")
pagebreak()

doc.add_heading("APPROVAL", level=1)
p("This is to certify that the final year project titled “Peonify: A Web-Based Floral "
  "E-Commerce and Order Management System” was carried out by [Student Full Name], "
  "Registration Number [Reg. No], under my supervision and guidance, and that it is now "
  "ready for submission and examination with my approval.")
p(); p("Supervisor: [Supervisor Name]")
p("Signature: ………………………………                Date: ………………………………")
pagebreak()

doc.add_heading("DEDICATION", level=1)
p("This work is dedicated to my beloved family, whose patience, sacrifices and constant "
  "encouragement have carried me through every stage of my studies, and to my friends and "
  "classmates who stood by me during the long days and nights this project required.")
pagebreak()

doc.add_heading("ACKNOWLEDGEMENTS", level=1)
p("First and foremost, I thank the Almighty God for the gift of life, health and strength "
  "throughout the period of my studies and of this project.")
p("I express my sincere gratitude to the University of Kigali, and in particular to the "
  "School of Computing and Information Technology, for providing the academic environment "
  "and resources within which this work was possible. I am especially indebted to my "
  "supervisor, [Supervisor Name], whose guidance, patience and constructive criticism "
  "shaped this project from an idea into a working system and a documented study.")
p("I also thank the flower entrepreneurs who generously shared their daily working "
  "experience during the requirements-gathering phase; their openness about the practical "
  "difficulties of selling flowers through informal channels grounded this project in a "
  "real problem. Finally, I thank my family and friends for their unconditional support.")
pagebreak()

# ╔══════════════════════════ TOC & LISTS ══════════════════════════╗
doc.add_heading("TABLE OF CONTENTS", level=1)
toc_field()
pagebreak()
LOF_HEAD = doc.add_heading("LIST OF FIGURES", level=1)
LOF_PAR = p("(generated below)")
doc.add_heading("LIST OF TABLES", level=1)
LOT_PAR = p("(generated below)")
pagebreak()
doc.add_heading("LIST OF ABBREVIATIONS AND ACRONYMS", level=1)
for ab in ["AJAX – Asynchronous JavaScript And XML",
           "API – Application Programming Interface",
           "CDN – Content Delivery Network",
           "CRUD – Create, Read, Update, Delete",
           "CSRF – Cross-Site Request Forgery",
           "CSS – Cascading Style Sheets",
           "ERD – Entity Relationship Diagram",
           "FR – Functional Requirement",
           "HTML – HyperText Markup Language",
           "HTTP – HyperText Transfer Protocol",
           "ICT – Information and Communication Technology",
           "IT – Information Technology",
           "JS – JavaScript",
           "JSON – JavaScript Object Notation",
           "NFR – Non-Functional Requirement",
           "OOAD – Object-Oriented Analysis and Design",
           "PDO – PHP Data Objects",
           "PHP – PHP: Hypertext Preprocessor",
           "SQL – Structured Query Language",
           "UAT – User Acceptance Testing",
           "UI – User Interface",
           "UML – Unified Modeling Language",
           "UoK – University of Kigali",
           "XAMPP – Cross-platform Apache MySQL PHP Perl"]:
    doc.add_paragraph(ab, style="List Bullet")
pagebreak()

# ╔══════════════════════════ CHAPTER 1 ══════════════════════════╗
doc.add_heading("CHAPTER 1", level=1)
doc.add_heading("GENERAL INTRODUCTION", level=1)

doc.add_heading("1.1 Introduction", level=2)
p("In today’s digital economy, businesses of every size are expected to serve their "
  "customers online. Electronic commerce has evolved from a competitive advantage into a "
  "baseline expectation: customers assume that they can browse a catalogue at any hour, "
  "compare prices transparently, pay securely, choose when a purchase will reach them, and "
  "follow the progress of their order without making a phone call. Enterprises that cannot "
  "offer this experience increasingly lose customers to those that can, regardless of the "
  "quality of the underlying product.")
p("This expectation now reaches even the smallest businesses. A neighbourhood florist "
  "competes, in the mind of the customer, with the polished ordering experience of "
  "international platforms. Yet the technical and financial barriers that once made such "
  "an experience unattainable for a micro-enterprise have largely disappeared: open-source "
  "server software, free relational databases and modern browsers make it possible to "
  "deliver professional commerce infrastructure at zero licence cost. This project applies "
  "that opportunity to a concrete case: it presents the analysis, design, implementation "
  "and testing of Peonify, a web-based floral e-commerce and order management system "
  "developed for a small flower boutique, using PHP, MySQL, HTML, CSS and vanilla "
  "JavaScript on the XAMPP platform.")

doc.add_heading("1.2 Background of the Study", level=2)
p("Flowers occupy a distinctive position in retail commerce. They are purchased primarily "
  "for emotional occasions — anniversaries, weddings, graduations, apologies, condolences "
  "— which makes reliability of delivery at least as important as the product itself. They "
  "are also highly perishable: an arrangement assembled in the morning must reach its "
  "recipient within hours to arrive at its best. Together these characteristics make "
  "timing and trust the two currencies of the floral trade.")
p("Globally, floral retail has moved decisively online. The flower delivery service market "
  "was estimated at USD 7.6 billion in 2024 and is projected to reach USD 11.27 billion by "
  "2030, growing at approximately seven percent per year (Grand View Research, 2024). "
  "Industry analyses report that online channels already account for roughly one third of "
  "all flower sales and that this share continues to grow annually (BusinessDojo, 2025). "
  "Established platforms such as 1-800-Flowers in the United States, Bloom & Wild in "
  "Europe and Floward in the Middle East have demonstrated that customers will confidently "
  "order premium, perishable arrangements from a screen when the experience communicates "
  "trust: authentic photography, transparent prices, precise delivery windows and visible "
  "order progress.")
p("In Rwanda, the environment for such systems is increasingly favourable. The country has "
  "made sustained investments in digital infrastructure and digital literacy, and a "
  "growing share of the population engages with online services primarily through mobile "
  "devices. Small businesses have responded by adopting the tools most immediately "
  "available to them: social-media platforms. Micro-entrepreneurs worldwide use Instagram, "
  "WhatsApp and similar applications as virtual storefronts because they are free and "
  "familiar (Mastercard Strive, 2024). The floral boutique studied in this project is "
  "typical of this pattern: talented at its craft, visible on social media, but entirely "
  "dependent on informal, manual processes for everything that happens after a customer "
  "says “I would like to order.”")
p("Studies of chat-based commerce consistently document the consequences of this "
  "dependency: fragmented customer journeys, manual order management, unstructured "
  "payments and the absence of any delivery discipline (PwC, 2024). Consumer trust is a "
  "further constraint — nearly half of surveyed consumers worry that purchases made "
  "through social media will not be protected or refunded if something goes wrong (Tidio, "
  "2024). The gap between what small florists can offer and what customers have learned "
  "to expect is precisely the gap this project addresses.")

doc.add_heading("1.3 Problem Statement", level=2)
p("The floral boutique that forms the case study of this project currently manages its "
  "entire commercial operation through informal channels. Product photographs are posted "
  "to social media; interested buyers send private messages; prices, delivery locations "
  "and delivery times are negotiated conversation by conversation; payment is collected in "
  "cash on delivery or through person-to-person mobile transfers; and the fulfilment of "
  "each order depends on the owner remembering it. This way of working is time-consuming "
  "and error-prone, and it produces no permanent record of what was sold, to whom, for "
  "how much, or whether it was ever delivered.")
p("As the volume of orders grows, the weaknesses of the manual process compound. Orders "
  "are forgotten or confused with one another; disputed prices cannot be resolved because "
  "no authoritative record exists; regular customers must reintroduce themselves with "
  "every purchase; and the owner cannot answer even elementary business questions such as "
  "how much revenue the boutique earned last month or which arrangements sell best. In "
  "the absence of a system, the ceiling of the business is the memory and availability of "
  "one person.")
doc.add_heading("1.3.1 Specific Problems", level=3)
bullets([
    "Ephemeral visibility: social-media posts disappear from followers’ feeds within "
    "hours, so the catalogue must be re-advertised continuously and has no permanent, "
    "searchable home.",
    "Informal ordering: orders captured in private chat threads suffer wrong addresses, "
    "missed dates and disputed prices, and can be neither tracked nor audited.",
    "Unstructured payments: cash and person-to-person transfers offer no receipts, no "
    "reconciliation against orders and no protection for either party.",
    "Absent customer relationships: without accounts, the business cannot notify "
    "customers of order progress, remind them of upcoming deliveries, or accumulate the "
    "verified feedback that builds trust for future buyers.",
    "No business intelligence: the owner has no view of revenue over time, of the order "
    "pipeline, or of product performance on which to base stocking and pricing decisions.",
])

doc.add_heading("1.4 Motivation of the Study", level=2)
p("The decision to develop Peonify is driven by the conviction that the benefits of "
  "digital commerce should not be reserved for large enterprises. The specific motivations "
  "of the study are threefold.")
p("To the University of Kigali, this study aligns with the university’s mission of "
  "producing graduates who apply information technology to solve practical problems in "
  "the surrounding community. It demonstrates the complete software engineering lifecycle "
  "— analysis, design, implementation, testing and documentation — applied to a genuine "
  "small-business problem rather than an abstract exercise.")
p("To the flower entrepreneur, the study delivers a working, zero-licence-cost system "
  "that professionalises daily operations: a permanent catalogue, verifiable payment "
  "records, disciplined deliveries, lasting customer relationships and actionable "
  "analytics — all operable without technical skills and installable by copying a single "
  "folder onto a Windows computer running XAMPP.")
p("To the student researcher, the project offered the opportunity to confront the "
  "realities of production software: security hardening, data integrity, usability for "
  "non-technical users, and the discipline of testing every user journey rather than only "
  "the successful path.")

doc.add_heading("1.5 Objectives of the Study", level=2)
doc.add_heading("1.5.1 General Objective", level=3)
p("The general objective of this study is to design and develop a web-based floral "
  "e-commerce and order management system that enables a flower entrepreneur to sell "
  "premium arrangements online with recorded payments, scheduled deliveries and "
  "transparent order progress, while remaining simple enough for a non-technical owner "
  "to install and operate end to end.")
doc.add_heading("1.5.2 Specific Objectives", level=3)
numbered([
    "To analyse the current, manual, social-media-based selling process of a small floral "
    "business and to model its shortcomings.",
    "To design a premium, responsive storefront — including an interactive bouquet "
    "builder with a live photographic preview — that converts visitors into buyers.",
    "To implement secure account management with distinct customer and administrator "
    "roles, using bcrypt password hashing, PHP sessions and anti-CSRF protection.",
    "To implement a checkout and payment-recording module that recomputes order totals on "
    "the server, marks orders as paid and issues unique, traceable order references.",
    "To implement a deliberately simple paid-to-delivered fulfilment workflow with "
    "in-application notifications and automatic delivery reminders for both parties.",
    "To provide the administrator with analytics — a thirty-day revenue chart, the order "
    "pipeline and top-selling products — together with complete catalogue control over "
    "products, categories, collections and promotional discounts.",
    "To test the developed system functionally across every user journey and document "
    "the results.",
])

doc.add_heading("1.6 Scope of the Study", level=2)
doc.add_heading("1.6.1 Content Scope", level=3)
p("The system covers the complete commercial cycle of a single floral boutique: the "
  "public storefront and merchandising; the bouquet builder; customer registration, "
  "authentication and profiles; the cart and checkout with delivery scheduling; payment "
  "recording with unique order references; order fulfilment and confirmation; "
  "notifications and delivery reminders; product feedback and its moderation; customer "
  "support messaging; and the administrator dashboard with analytics and catalogue "
  "management.")
doc.add_heading("1.6.2 Geographical Scope", level=3)
p("The case study targets a floral boutique operating in Kigali, Rwanda. Kigali was "
  "chosen because it concentrates the customer base most likely to order flowers online, "
  "enjoys the country’s highest levels of internet penetration and digital literacy, and "
  "hosts the events — weddings, conferences, graduations — that drive premium floral "
  "demand.")
doc.add_heading("1.6.3 Technological Scope", level=3)
p("The system is implemented for the XAMPP platform: Apache as the web server, PHP 8 as "
  "the server-side language, MySQL 8 as the relational database, and HTML5, CSS3 and "
  "vanilla JavaScript on the client. No JavaScript frameworks and no build tooling are "
  "used, so the application runs directly from the web server’s document root.")
doc.add_heading("1.6.4 Out of Scope", level=3)
bullets([
    "Multi-vendor marketplace functionality — the system serves a single merchant by design.",
    "Native mobile applications — the responsive web interface serves mobile browsers.",
    "Courier fleet management and third-party logistics integration.",
    "Subscription or recurring billing.",
    "Integration with an external online payment gateway, which is identified as future work.",
])

doc.add_heading("1.7 Methods and Techniques Used in the Study", level=2)
doc.add_heading("1.7.1 Data Collection Techniques", level=3)
p("Three complementary techniques were used to establish the requirements of the system. "
  "Documentation review examined how established floral platforms — 1-800-Flowers, "
  "Floward and Bloom & Wild — and general marketplaces present products, schedule "
  "deliveries and communicate order progress, in order to identify the interaction "
  "patterns customers already understand. Direct observation followed the day-to-day "
  "selling activity of small florists on Instagram and WhatsApp, recording how orders are "
  "initiated, negotiated and too often lost. Finally, informal interviews with the target "
  "entrepreneur established the practical constraints of the business: no technical "
  "staff, a single computer running Windows, and no budget for software licences or "
  "monthly platform fees.")
doc.add_heading("1.7.2 System Development Methodology", level=3)
p("Development followed the Agile iterative model. The system was built in vertical "
  "slices — database, server logic, user interface and test for one feature at a time — "
  "and each slice was demonstrated and refined before the next began. This approach "
  "proved essential: several design decisions, most notably the simplification of the "
  "fulfilment pipeline into a single paid-to-delivered transition and the addition of "
  "customer delivery confirmation, emerged directly from iteration feedback rather than "
  "from the initial design.")
doc.add_heading("1.7.3 Analysis and Design Techniques", level=3)
p("Object-oriented analysis and design (OOAD) was used to model the system. Use case "
  "modelling captured functional requirements from the perspective of each actor; class "
  "modelling identified the principal entities and their responsibilities; sequence "
  "modelling documented the checkout interaction; and relational design mapped the model "
  "onto a normalised MySQL schema. The Unified Modeling Language (UML) notation is used "
  "throughout Chapter 3.")

doc.add_heading("1.8 Significance of the Study", level=2)
p("For the boutique, the significance is immediate and operational: a permanent, "
  "searchable catalogue replaces ephemeral posts; recorded payments with unique "
  "references replace unverifiable transfers; scheduled, reminded deliveries replace "
  "promises made from memory; and a dashboard finally answers the questions every owner "
  "asks — what is selling, what is pending, and how is revenue moving.")
p("For the wider small-business community, the project provides a replicable template. "
  "Because the entire stack is open-source and the deployment procedure is a folder copy, "
  "the same architecture can serve bakeries, tailors, artisans and any other single-"
  "merchant business that today depends on chat-based selling.")
p("For the academic community, the project documents a complete, security-conscious "
  "implementation of a real commerce system on the classic PHP/MySQL stack that dominates "
  "small-business hosting worldwide, demonstrating that rigorous engineering practice is "
  "compatible with the simplest and most accessible tools.")

doc.add_heading("1.9 Expected Results", level=2)
bullets([
    "A permanent, professionally merchandised online boutique replacing ephemeral "
    "social-media posts.",
    "Recorded, reconcilable payments with unique order references and automatic "
    "notifications to both parties.",
    "A complete and auditable order history, with delivery reminders sent twenty-four "
    "hours and five hours before each delivery window.",
    "Verified customer feedback displayed on products and on the landing page, with "
    "administrative moderation.",
    "Actionable analytics for the entrepreneur: daily revenue over thirty days, the "
    "order pipeline and top-selling arrangements.",
    "A deployment procedure achievable by a non-developer: install XAMPP, copy one "
    "folder, open the browser.",
])

doc.add_heading("1.10 Definition of Key Terms", level=2)
bullets([
    "E-commerce: the buying and selling of goods or services over the internet, "
    "including the payment and fulfilment processes that complete each transaction.",
    "Order management: the recording, tracking and completion of customer orders from "
    "placement through delivery, with an auditable history of every state change.",
    "Storefront: the public, customer-facing portion of an e-commerce system where "
    "products are presented, discovered and added to a cart.",
    "Checkout: the process in which a customer provides delivery details, the system "
    "computes the authoritative total, payment is recorded and an order is created.",
    "Fulfilment: the operational completion of a paid order, ending with delivery to "
    "the customer and confirmation within the system.",
    "Session: the server-maintained association between a signed-in user and their "
    "subsequent requests, carried by a protected browser cookie.",
    "Seeding: the automatic insertion of initial data — accounts, taxonomy and "
    "demonstration products — when a system is installed.",
])

doc.add_heading("1.11 Organization of the Report", level=2)
p("The remainder of this report is organised as follows. Chapter two analyses the "
  "existing manual system, models its process, articulates its problems and derives the "
  "functional and non-functional requirements of the proposed solution. Chapter three "
  "presents the requirements analysis and design of the new system using UML models, the "
  "database schema and the system architecture. Chapter four discusses the technical "
  "implementation — the technologies used and their justification, the realisation of "
  "each module, the security measures applied — and presents the developed system through "
  "authentic screen captures together with the testing performed. Chapter five concludes "
  "the study and offers recommendations and directions for future work.")
pagebreak()

# ╔══════════════════════════ CHAPTER 2 ══════════════════════════╗
doc.add_heading("CHAPTER 2", level=1)
doc.add_heading("ANALYSIS OF THE EXISTING SYSTEM", level=1)

doc.add_heading("2.1 Introduction", level=2)
p("Before a new system can be designed responsibly, the system it replaces must be "
  "understood in detail — not only its mechanics but the reasons it persists and the "
  "precise points at which it fails. This chapter describes the environment and history "
  "of the boutique’s current way of working, models the existing manual process, "
  "analyses its problems, presents the proposed solution, and derives the functional and "
  "non-functional requirements that guided the design in Chapter 3.")

doc.add_heading("2.2 Description of the Existing System Environment", level=2)
doc.add_heading("2.2.1 Historical Background", level=3)
p("The boutique began as a home-based flower business promoted through the owner’s "
  "personal social-media accounts. Early orders came from friends and acquaintances, for "
  "whom informality was natural: a message, an agreed price, a delivery arranged in "
  "person. As word spread, the owner adopted a business Instagram profile and WhatsApp "
  "Business for order taking, and order volume grew beyond the circle of personal "
  "acquaintances. The tools, however, did not change. All records — orders, payments, "
  "delivery promises — remained where they were born: inside chat histories and, "
  "occasionally, a personal notebook. No computerised system of any kind has ever been "
  "used in the business.")
doc.add_heading("2.2.2 Current Operating Context", level=3)
p("Today the boutique sources fresh stems from local growers, assembles arrangements in "
  "a small workshop, and sells through two channels: walk-in customers, and social-media "
  "direct messages that now account for the majority of orders. Deliveries within Kigali "
  "are made personally by the owner or by informal courier arrangements agreed order by "
  "order. The owner operates alone, with occasional family help during peak periods such "
  "as Valentine’s season and graduation weeks — precisely the periods in which the "
  "manual process is most likely to fail.")

doc.add_heading("2.3 Description of the Existing System", level=2)
p("The existing “system” is a sequence of manual activities threaded through chat "
  "applications. A typical order proceeds as follows. The owner photographs finished "
  "arrangements and posts them to Instagram with an invitation to order by direct "
  "message. An interested buyer writes privately, often beginning a negotiation over "
  "price, stem selection, wrapping, delivery place and time. When agreement is reached — "
  "sometimes across dozens of messages spread over hours — the buyer either promises "
  "cash on delivery or sends a mobile-money transfer, of which the only record is the "
  "transfer message itself. The owner then writes the delivery into memory or a "
  "notebook, prepares the arrangement on the agreed morning, and delivers it. If the "
  "buyer later wishes to order again, the entire conversation begins from zero; nothing "
  "about them — address, preferences, history — has been retained anywhere searchable.")
p("Feedback, when it exists, arrives as a private thank-you message seen only by the "
  "owner. It builds no public trust and influences no future buyer. Equally, complaints "
  "arrive privately and leave no trace that could improve the process.")

doc.add_heading("2.4 Analysis of the Existing System", level=2)
doc.add_heading("2.4.1 Modeling of the Existing System", level=3)
p("Figure 1 models the existing process as a flow of five manual activities. Beneath "
  "each activity, the analysis notes the characteristic failure it introduces. The model "
  "makes visible what the participants experience daily: every stage depends on human "
  "memory and goodwill, and no stage produces a durable, verifiable record.")
diagram(dg.existing_process, "Model of the existing manual selling process")
p("Reading the model from left to right: visibility exists only while a post is recent; "
  "the order enters a private thread indistinguishable from dozens of others; the "
  "commercial terms are improvised rather than standing; the payment leaves no "
  "reconcilable record; and the delivery depends on recollection. A failure at any stage "
  "is invisible until a customer complains.")

doc.add_heading("2.5 Problems of the Existing System", level=2)
p("The analysis identifies six systemic problems, each traceable to a stage of the model "
  "above.")
bullets([
    "No permanent catalogue: every product must be re-posted to remain visible, and a "
    "buyer who saw an arrangement last week has no way to find it today.",
    "Scattered, unauditable orders: commitments live in chat threads where they are "
    "easily forgotten, confused between customers, or lost when a device changes.",
    "Unverifiable payments: cash and person-to-person transfers cannot be reconciled "
    "against orders; disputes reduce to one person’s word against another’s.",
    "Undisciplined delivery: with no scheduling and no reminders, deliveries are missed "
    "or late precisely during high-demand periods, damaging the trust on which an "
    "emotional purchase depends.",
    "No accumulated customer relationships: addresses, preferences and history evaporate "
    "with each conversation, making loyalty invisible and repeat purchase needlessly "
    "difficult.",
    "No reporting of any kind: the owner cannot state monthly revenue, count pending "
    "orders, or identify the best-selling arrangement other than by impression.",
])

doc.add_heading("2.6 Review of Similar Systems", level=2)
p("Before designing a replacement, existing ways of selling flowers online were "
  "reviewed to identify both the interaction patterns customers already understand and "
  "the gaps that justify a purpose-built system for a single boutique.")
doc.add_heading("2.6.1 Social-Media Selling (Instagram and WhatsApp Business)", level=3)
p("Social platforms provide instant reach at no cost and remain the de facto standard "
  "for small florists — indeed, they are the existing system of this study. Their "
  "strengths end at visibility: they provide no catalogue structure, no ordering, no "
  "payment records and no delivery discipline, and the platform’s algorithm rather "
  "than the merchant decides who sees each product. They demonstrate the audience but "
  "not the infrastructure.")
doc.add_heading("2.6.2 General Marketplaces (Alibaba, Kikuu, Jumia)", level=3)
p("General marketplaces contribute the commerce mechanics customers already know: "
  "structured listings, carts, discount badges, best-seller sections and buyer "
  "reviews. Peonify deliberately borrows these patterns. As a home for a premium "
  "perishable product, however, marketplaces fall short: a bouquet requiring same-day "
  "delivery within a chosen time window does not fit commodity logistics, the product "
  "drowns among unrelated goods, and the merchant surrenders both brand identity and "
  "the customer relationship to the platform.")
doc.add_heading("2.6.3 Hosted Store Builders (Shopify and Similar)", level=3)
p("Hosted builders give a merchant a branded store with themes and plugins, and prove "
  "that non-developers can run online shops. Their limits for this case are economic "
  "and functional: recurring monthly fees are significant for a micro-enterprise; "
  "florist-specific needs — delivery windows, bouquet configuration, delivery "
  "reminders — require additional paid applications; and the administration surface "
  "remains generic and complex relative to the owner’s actual daily tasks.")
doc.add_heading("2.6.4 Dedicated Floral Platforms (1-800-Flowers, Floward, Bloom & Wild)", level=3)
p("The dedicated platforms represent the customer-experience benchmark: authentic "
  "photography, curated collections, precise delivery scheduling and visible order "
  "progress. Floward, for example, built its regional leadership on reliable same-day "
  "delivery within tight windows. These platforms are proprietary businesses rather "
  "than software an independent florist can adopt; joining them, where possible at "
  "all, means surrendering margin and brand. Peonify adopts their proven experience "
  "patterns in software the entrepreneur owns outright.")
doc.add_heading("2.6.5 Comparative Summary", level=3)
table(["System", "Strengths", "Weaknesses for this case"],
    [["Social-media selling", "Free; instant reach; visual medium suits flowers", "No orders, payments, scheduling or records; algorithm-dependent visibility"],
     ["General marketplaces", "Familiar commerce mechanics; integrated payments; reviews", "Commodity logistics unfit for perishables; brand and customer data surrendered"],
     ["Hosted store builders", "Branded store without coding; plugin ecosystem", "Recurring fees; florist features cost extra; complex generic administration"],
     ["Dedicated floral platforms", "Benchmark delivery experience and trust", "Closed businesses, not adoptable software; margin and brand loss"],
     ["Peonify (proposed)", "Purpose-built florist workflows; owner-operated; zero licence cost; self-installing", "Single merchant by design; placeholder imagery until owner photography is added"]],
    cap="Comparison of existing approaches with the proposed system", widths=[1.5, 2.4, 2.5])

doc.add_heading("2.7 Proposed Solution", level=2)
p("The proposed solution is Peonify, a self-contained web-based e-commerce and order "
  "management system installed on the XAMPP platform. Peonify addresses each identified "
  "problem directly. A permanent, searchable storefront with professional merchandising "
  "replaces ephemeral posts. Account-based ordering through a structured checkout — with "
  "recipient details, a delivery date and a selectable time window — replaces chat "
  "negotiation. Payment is recorded at checkout with the order total recomputed on the "
  "server and a unique order reference issued for every purchase, replacing unverifiable "
  "transfers. A deliberately simple fulfilment workflow, automatic notifications on both "
  "sides and delivery reminders twenty-four hours and five hours before the chosen "
  "window replace memory-based delivery. Customer accounts retain addresses and order "
  "history, and verified customers may rate and review products, accumulating public "
  "trust. Finally, an administrator dashboard presents revenue analytics, the order "
  "pipeline and catalogue management in plain language designed for a non-technical "
  "owner.")
p("Critically for this case, the system is engineered for autonomous deployment: on its "
  "first visit it creates its own database, all tables and a demonstration catalogue, so "
  "that installation on a Windows computer consists of installing XAMPP and copying one "
  "folder into the web server’s document root.")

doc.add_heading("2.8 System Requirements", level=2)
p("The requirements below were derived from the problem analysis and validated against "
  "the entrepreneur’s working needs. They are stated as testable capabilities; the test "
  "results appear in Chapter 4.")
doc.add_heading("2.8.1 Functional Requirements", level=3)
table(["Code", "Requirement", "Description"],
    [["FR1", "User registration and authentication", "The system shall allow users to register and sign in with an email address and password. Passwords shall be stored only as bcrypt hashes; sessions shall use HTTP-only cookies; the administrator account shall be seeded automatically on first run."],
     ["FR2", "Role-based access control", "The system shall distinguish customers from the administrator. Administrators shall not be able to purchase; customers shall not be able to reach management functions; each role shall be directed to its own dashboard."],
     ["FR3", "Product catalogue management", "The administrator shall create, edit and delete products — with photo upload, price, category, collection, stock status and an optional promotional discount — through a modal form, with changes visible in the storefront immediately."],
     ["FR4", "Storefront and discovery", "Customers shall browse with text search, category chips, a collection filter, a price-range filter, a sale-only filter, sorting and pagination. The landing page shall merchandise best sellers, new arrivals, flash sales, collections and customer testimonials."],
     ["FR5", "Bouquet builder", "The system shall offer a four-step configurator (size, focal flower, foliage, packaging) with a live photographic preview and a per-step price breakdown; the chosen configuration shall be stored with the order line."],
     ["FR6", "Cart and checkout", "The cart shall persist in the browser. Checkout shall require an account, pre-fill recipient details from the profile, and capture the delivery address, date, time window and an optional gift note."],
     ["FR7", "Payment recording", "At checkout the system shall recompute every price on the server from the database, record the payment, mark the order as paid, and issue a unique human-readable reference (e.g., PNY-A1B2C3)."],
     ["FR8", "Fulfilment workflow", "Orders shall move from paid to delivered through a single action: the administrator pressing Deliver (with a confirmation dialog) or the receiving customer pressing “I received it”. Every transition shall be logged as a timestamped event."],
     ["FR9", "Notifications and reminders", "Both roles shall receive in-application notifications with unread badges and per-item and mark-all read controls. Customers shall be reminded twenty-four hours and five hours before their delivery window."],
     ["FR10", "Feedback and support", "Verified customers shall rate and review products (one review per customer per product, editable). The public contact form and the account support tab shall deliver messages to the administrator’s inbox. The administrator shall be able to remove inappropriate reviews."],
     ["FR11", "Analytics dashboard", "The administrator dashboard shall chart daily revenue over the last thirty days and the order pipeline, list top-selling products, and display headline statistics with compact money formatting."],
     ["FR12", "Profile management", "Both roles shall manage a profile: a click-to-upload photograph, name, phone, and a saved delivery address and city used to pre-fill checkout; password change shall require the current password and a matching confirmation."]],
    cap="Functional requirements of the proposed system", widths=[0.55, 1.55, 4.3])
doc.add_heading("2.8.2 Non-Functional Requirements", level=3)
table(["Code", "Requirement", "Description"],
    [["NFR1", "Usability", "All interfaces shall use plain language; the administrator dashboard shall be operable by a non-developer; destructive actions shall require confirmation."],
     ["NFR2", "Security", "The system shall apply bcrypt hashing (cost 12), HTTP-only SameSite session cookies, anti-CSRF tokens on every form, prepared SQL statements exclusively, image-only validated uploads limited to five megabytes, and server-side price recomputation."],
     ["NFR3", "Performance", "Long lists shall be paginated; images shall load lazily; each request shall use a single pooled database connection; interface libraries shall be served from CDNs."],
     ["NFR4", "Portability", "The system shall run on a default XAMPP installation on Windows without configuration; the database shall create and seed itself on first run."],
     ["NFR5", "Reliability", "Order creation shall be transactional; totals shall always be recomputed on the server before an order is stored."],
     ["NFR6", "Maintainability", "The code shall be organised as one page-controller per screen with shared include files; schema migration and seeding shall be idempotent."],
     ["NFR7", "Responsiveness", "All screens, including the administrator dashboard, shall adapt to mobile viewports."],
     ["NFR8", "Availability", "The system shall operate wholly on the local server without external service dependencies for its core flows."]],
    cap="Non-functional requirements of the proposed system", widths=[0.6, 1.3, 4.5])

doc.add_heading("2.9 Feasibility of the Proposed Solution", level=2)
p("Technical feasibility: every component of the chosen stack — Apache, PHP 8, MySQL 8, "
  "HTML5, CSS3 and vanilla JavaScript — is mature, freely available and documented "
  "exhaustively; no experimental technology is required.")
p("Operational feasibility: the administrator interface was designed with, and validated "
  "against, a non-technical user. Every operational task — adding a product, delivering "
  "an order, answering a message — is a single screen with plain-language labels and "
  "confirmation dialogs.")
p("Economic feasibility: the entire software stack carries zero licence cost. The only "
  "recurring costs of a production deployment are a domain name and basic hosting, both "
  "of which are widely affordable to a functioning boutique.")
pagebreak()
# -*- coding: utf-8 -*- (concatenated after build_report_1.py)

# ╔══════════════════════════ CHAPTER 3 ══════════════════════════╗
doc.add_heading("CHAPTER 3", level=1)
doc.add_heading("REQUIREMENTS ANALYSIS AND DESIGN OF THE NEW SYSTEM", level=1)

doc.add_heading("3.1 Introduction", level=2)
p("This chapter translates the requirements established in Chapter 2 into a concrete "
  "design. It first explains the analysis methodology and the development process "
  "followed, then presents the system through a sequence of UML models — the use case "
  "model with detailed use case descriptions, the class model, the checkout sequence and "
  "the order life cycle — before describing the relational database design and the "
  "overall system architecture. All diagrams in this chapter are drawn as native, "
  "editable drawings; the database schema is presented as captured from the phpMyAdmin "
  "Designer of the live system.")

doc.add_heading("3.2 Object-Oriented Methodology", level=2)
p("Object-Oriented Analysis and Design was selected because the domain decomposes "
  "naturally into cooperating entities with clear responsibilities: users who "
  "authenticate and maintain profiles; products that carry pricing and presentation; "
  "orders that bind customers, items and delivery details together; and services that "
  "cut across entities, such as notification delivery. During analysis, use case "
  "modelling captured what each actor must be able to accomplish without prescribing "
  "how. During design, class modelling assigned each responsibility to exactly one "
  "entity, and sequence modelling verified that the entities could collaborate to "
  "fulfil the most important scenario — checkout — without gaps or ambiguity. The "
  "resulting model maps directly onto the relational schema of Section 3.9.")

doc.add_heading("3.3 Software Development Process", level=2)
p("The Agile iterative model governed construction. Each iteration selected a coherent "
  "slice of functionality, implemented it across all three tiers, and ended with "
  "functional testing and a review against the entrepreneur’s expectations. Two design "
  "corrections illustrate the value of this discipline. First, an early design carried a "
  "six-stage fulfilment pipeline; owner review showed that the intermediate stages "
  "created work without informing anyone, and the pipeline was simplified to the single "
  "paid-to-delivered transition that shipped. Second, delivery confirmation was initially "
  "an administrator-only action; observation showed that customers often receive flowers "
  "before the administrator updates the system, so a customer-side “I received it” "
  "confirmation was added. Figure 2 presents the four-month project timeline.")
diagram(dg.gantt, "Project timeline — Gantt chart across four months")

doc.add_heading("3.4 Use Case Diagram", level=2)
p("Figure 3 presents the use case model. Three human actors interact with the system. "
  "The Guest browses the catalogue, explores the bouquet builder and may register or "
  "log in. The Customer, once authenticated, gains the commercial capabilities: checkout "
  "and payment, delivery confirmation, product feedback, profile management, support "
  "contact and notifications. The Administrator operates the boutique: product and "
  "catalogue management, promotional discounts, order delivery, analytics, feedback "
  "moderation, the support inbox and the activity audit. Guests inherit nothing that "
  "requires identity; customers inherit all guest capabilities.")
diagram(dg.usecase, "Use case diagram of the Peonify web system")

doc.add_heading("3.5 Use Case Descriptions", level=2)
p("The eight most significant use cases are specified below in the standard tabular "
  "form: actors, preconditions, the main flow and postconditions.")
UC = [
    ("Register", "Guest",
     "The visitor has no account.",
     "1. The guest opens the registration page. 2. The guest enters full name, email, "
     "password and password confirmation, with visibility toggles available. 3. The "
     "system validates the email format, the minimum password length of eight "
     "characters, and the match between password and confirmation. 4. The system stores "
     "the account with a bcrypt hash and starts a session. 5. The new customer is "
     "greeted on their account dashboard.",
     "A customer account exists and is signed in."),
    ("Login", "Guest (becoming Customer or Administrator)",
     "A registered account exists.",
     "1. The user opens the login page. 2. The user submits email and password. 3. The "
     "system verifies the bcrypt hash, regenerates the session identifier and stores "
     "the user identity in the session. 4. The system redirects to the dashboard "
     "matching the account’s role.",
     "An authenticated session exists; the notification badge reflects unread items."),
    ("Browse and Search the Shop", "Guest / Customer",
     "None.",
     "1. The visitor opens the shop. 2. The visitor optionally narrows the view with "
     "the text search, a category chip, the collection filter, a price range or the "
     "sale-only filter, and chooses a sort order. 3. The system applies all criteria, "
     "recomputes discounted prices, and presents a paginated grid with badges for new "
     "and discounted arrangements.",
     "The visitor views the matching subset of the catalogue."),
    ("Build a Custom Bouquet", "Guest / Customer",
     "None.",
     "1. The visitor opens the bouquet builder. 2. The visitor selects one option in "
     "each of four steps: size, focal flower, foliage and packaging. 3. With each "
     "selection the system updates a live photographic preview and a per-step price "
     "breakdown. 4. When all four steps are chosen, the visitor adds the configured "
     "bouquet to the cart.",
     "A custom bouquet with its configuration and computed price is in the cart."),
    ("Checkout and Pay", "Customer",
     "The customer is signed in and the cart contains at least one item.",
     "1. The customer opens checkout; recipient details are pre-filled from the "
     "profile. 2. The customer chooses on-demand or a future date, selects a delivery "
     "window, and optionally adds a gift note. 3. On submission the system verifies the "
     "anti-CSRF token, recomputes every price from the database, creates the order and "
     "its items inside a transaction, records payment, marks the order paid and issues "
     "a unique reference. 4. The system notifies the customer and the administrator "
     "and clears the cart.",
     "A paid order with an auditable event trail exists."),
    ("Deliver Order", "Administrator",
     "A paid order exists.",
     "1. The administrator opens Orders and locates the order by search or status "
     "filter. 2. The administrator presses Deliver and confirms in the dialog. 3. The "
     "system marks the order delivered, records the event and notifies the customer. "
     "Alternate flow: the receiving customer presses “I received it” on their own "
     "order, with the administrator notified instead.",
     "The order is completed and the customer may leave feedback."),
    ("Rate and Review a Product", "Customer",
     "The customer is signed in.",
     "1. The customer opens a product page and selects a star rating with an optional "
     "comment. 2. The system stores one review per customer per product, replacing any "
     "earlier review by the same customer. 3. The administrator is notified; the "
     "review appears on the product page and, if highly rated, among the landing-page "
     "testimonials.",
     "A verified review is published and available for moderation."),
    ("Receive Notifications", "Customer / Administrator",
     "The user is signed in.",
     "1. The system creates notifications when orders are placed and delivered, when "
     "messages and reviews arrive, and as delivery reminders twenty-four hours and "
     "five hours before each window. 2. The user opens the notification list from the "
     "badge and marks single items or all items as read.",
     "The unread badge reflects the remaining unread notifications."),
]
for name, actor, pre, flow, post in UC:
    table(["Field", "Description"],
        [["Use case name", name], ["Actor(s)", actor], ["Precondition", pre],
         ["Main flow", flow], ["Postcondition", post]],
        cap=f"Use case description — {name}", widths=[1.35, 5.05])

doc.add_heading("3.6 Class Diagram", level=2)
p("Figure 4 presents the class model. User carries identity, role and the saved "
  "delivery details that pre-fill checkout. Product owns presentation and pricing, "
  "including the promotional discount from which its effective price is derived. Order "
  "aggregates OrderItems — each of which snapshots the product name and unit price at "
  "purchase time, so later catalogue changes cannot rewrite history — and exposes the "
  "two operations of the fulfilment workflow. Review associates a customer with a "
  "product under a uniqueness constraint. NotificationService is the cross-cutting "
  "collaborator through which every significant event reaches its audience, including "
  "the scheduled delivery reminders.")
diagram(dg.classes, "Class diagram of the principal entities")

doc.add_heading("3.7 Sequence Diagram — Checkout", level=2)
p("Figure 5 verifies the collaboration for the system’s most important scenario. Note "
  "two properties of the design that the diagram makes explicit. First, the browser’s "
  "cart contents are treated as a proposal only: the server recomputes every price from "
  "the database before anything is stored, so a manipulated client cannot alter what it "
  "pays. Second, order creation is transactional: the order, its items and its first "
  "event either all exist or none do.")
diagram(dg.sequence, "Sequence diagram — checkout and order creation")

doc.add_heading("3.8 Order Life Cycle", level=2)
p("Figure 6 presents the deliberately minimal state model that emerged from iteration "
  "with the owner. An order is paid the moment checkout completes and delivered through "
  "a single action by either party. Every transition is recorded in the order_events "
  "table, which doubles as the administrator’s audit trail.")
diagram(dg.orderflow, "Order life cycle state model")

doc.add_heading("3.9 Database Design", level=2)
p("The relational schema was designed in third normal form: every attribute depends on "
  "the key of its table, catalogue taxonomy is factored into its own tables, and all "
  "cross-entity references are enforced with foreign keys. Deletion behaviour is chosen "
  "per relationship — order items survive product deletion with their snapshot data "
  "(SET NULL), while notifications and reviews follow their owning user (CASCADE). "
  "Figure 7 shows the schema as captured from the phpMyAdmin Designer of the live "
  "peonify database.")
picture(DBSHOT, "Database schema diagram (phpMyAdmin Designer view of the live peonify database)")
p("The entities are specified below.")
ENTS = [
    ("users", "Customers and the administrator, distinguished by the role column.",
     [("id", "INT, PK, auto-increment", "Unique user identifier"),
      ("name / email", "VARCHAR; email UNIQUE", "Identity and sign-in credential"),
      ("password_hash", "VARCHAR(255)", "bcrypt hash; plaintext is never stored"),
      ("role", "ENUM('customer','admin')", "Access level"),
      ("avatar_url / phone", "VARCHAR", "Profile photograph and contact"),
      ("address / city", "VARCHAR", "Saved delivery location pre-filling checkout"),
      ("created_at", "TIMESTAMP", "Registration time")]),
    ("products", "The catalogue with pricing, discounting and imagery.",
     [("id / slug", "INT PK; VARCHAR UNIQUE", "Identifier and URL-safe name"),
      ("name / description", "VARCHAR / TEXT", "Presentation"),
      ("category / collection", "VARCHAR", "Taxonomy (managed in categories / collections)"),
      ("price_cents", "INT", "Price in cents — exact integer arithmetic"),
      ("discount_percent", "INT 0–90", "Promotional (flash-sale) discount"),
      ("image_url / in_stock", "VARCHAR / TINYINT", "Photograph and availability"),
      ("created_at", "TIMESTAMP", "Drives the “New” badge and newest-first sort")]),
    ("orders", "One row per purchase.",
     [("id / reference", "INT PK; VARCHAR UNIQUE", "Identifier and human-readable reference (PNY-XXXXXX)"),
      ("user_id", "INT, FK→users, SET NULL", "Purchasing customer"),
      ("customer_name / email / phone", "VARCHAR", "Recipient details"),
      ("address / delivery_date / delivery_window", "VARCHAR / DATE / VARCHAR", "Where and when to deliver"),
      ("order_type / gift_note", "VARCHAR / TEXT", "On-demand or scheduled; optional note"),
      ("total_cents / payment_ref", "INT / VARCHAR", "Server-computed total and payment record"),
      ("status", "ENUM paid | delivered (pending_payment reserved)", "Fulfilment state"),
      ("reminded_1d / reminded_5h", "TINYINT", "Reminder bookkeeping"),
      ("created_at", "TIMESTAMP", "Purchase time")]),
    ("order_items", "The lines of an order; snapshots survive catalogue changes.",
     [("id / order_id", "INT PK; FK→orders CASCADE", "Line and owning order"),
      ("product_id", "INT, FK→products, SET NULL", "Product, if catalogue-based"),
      ("name / quantity / unit_price_cents", "VARCHAR / INT / INT", "Snapshot at purchase time"),
      ("custom_config", "TEXT (JSON)", "Bouquet-builder configuration, when custom")]),
    ("order_events", "The audit trail of every status change.",
     [("id / order_id", "INT PK; FK→orders CASCADE", "Event and owning order"),
      ("status / note", "VARCHAR", "Transition and human-readable note"),
      ("created_at", "TIMESTAMP", "When the transition occurred")]),
    ("notifications", "In-application notifications for both roles.",
     [("id / user_id", "INT PK; FK→users CASCADE", "Notification and recipient"),
      ("title / body / link", "VARCHAR", "Content"),
      ("is_read", "TINYINT", "Read state; drives unread badges"),
      ("created_at", "TIMESTAMP", "Creation time")]),
    ("reviews", "Verified product feedback.",
     [("id", "INT PK", "Identifier"),
      ("user_id / product_id", "FKs, CASCADE; UNIQUE(user_id, product_id)", "One review per customer per product"),
      ("rating / comment", "TINYINT 1–5 / TEXT", "The feedback itself"),
      ("created_at", "TIMESTAMP", "Last submission time")]),
    ("messages", "Contact-form and support messages (administrator inbox).",
     [("id", "INT PK", "Identifier"),
      ("name / email / subject / body", "VARCHAR / TEXT", "Sender and content"),
      ("created_at", "TIMESTAMP", "Arrival time")]),
    ("categories & collections", "Administrator-managed catalogue taxonomy.",
     [("id / slug / name", "INT PK; VARCHAR UNIQUE; VARCHAR", "Identifier and labels"),
      ("description", "VARCHAR (collections only)", "Shown on the landing page")]),
    ("builder_options", "The bouquet builder’s selectable options.",
     [("id / step", "INT PK; ENUM size|focal|foliage|packaging", "Option and its step"),
      ("name / detail / price_cents", "VARCHAR / VARCHAR / INT", "Label, description and price contribution")]),
    ("settings", "Small key-value store for system bookkeeping.",
     [("k / v", "VARCHAR PK / VARCHAR", "Used to pace the reminder scheduler")]),
]
for ent, desc, cols in ENTS:
    doc.add_heading(f"Entity: {ent}", level=4)
    p(desc)
    table(["Attribute", "Type / Constraint", "Purpose"], cols,
          cap=f"Attributes of the {ent} entity", widths=[1.7, 2.2, 2.5])

doc.add_heading("3.10 System Architecture Design", level=2)
p("Figure 8 presents the three-tier architecture. The presentation tier runs entirely "
  "in the browser with no frameworks and no build step: semantic HTML5, a hand-written "
  "CSS3 theme, and vanilla JavaScript providing the cart, the bouquet preview, toasts, "
  "modal dialogs and confirmation dialogs, with Lucide supplying icons and Chart.js "
  "rendering the dashboard charts. The application tier consists of one PHP page "
  "controller per screen, sharing two include files: db.php, which owns the PDO "
  "connection and the idempotent migration and seeding logic, and functions.php, which "
  "provides authentication, CSRF protection, uploads, notifications and the reminder "
  "scheduler. The data tier is MySQL 8 with InnoDB tables and enforced foreign keys. "
  "The separation allows each tier to evolve independently: the storefront can be "
  "restyled without touching order logic, and the schema migrates itself without manual "
  "administration.")
diagram(dg.architecture, "Three-tier system architecture on the XAMPP platform")

doc.add_heading("3.11 Interface Design Principles", level=2)
p("The interface design pursued two goals that are usually in tension: a premium visual "
  "identity appropriate to a luxury product, and operational simplicity appropriate to "
  "non-technical users. The visual identity rests on a rose-and-ivory palette, a serif "
  "display face (Cormorant Garamond) paired with a clean sans-serif (Jost), authentic "
  "flower photography, and consistent iconography. Operational simplicity rests on "
  "plain-language labels, one task per screen, confirmation dialogs before destructive "
  "actions, toast feedback after every action, and pagination wherever content can "
  "grow. Chapter 4 presents the realised interfaces.")
pagebreak()
# -*- coding: utf-8 -*- (concatenated after build_report_2.py)

# ╔══════════════════════════ CHAPTER 4 ══════════════════════════╗
doc.add_heading("CHAPTER 4", level=1)
doc.add_heading("SYSTEM IMPLEMENTATION AND TESTING", level=1)

doc.add_heading("4.1 Introduction", level=2)
p("This chapter presents the technical realisation of Peonify. It describes the "
  "development environment and the technologies employed with their justification, "
  "explains how each module of the design was implemented, details the security "
  "measures applied, presents the developed system through authentic screen captures of "
  "the running application, and reports the testing performed against the requirements "
  "of Chapter 2.")

doc.add_heading("4.2 Development Environment", level=2)
doc.add_heading("4.2.1 Hardware Requirements", level=3)
table(["Component", "Minimum", "Used during development"],
    [["Processor", "Dual-core 1.6 GHz", "Quad-core 2.4 GHz"],
     ["Memory", "4 GB RAM", "8 GB RAM"],
     ["Storage", "2 GB free (XAMPP + system + images)", "SSD storage"],
     ["Display", "1366 × 768", "1920 × 1080"]],
    cap="Hardware requirements", widths=[1.6, 2.4, 2.4])
doc.add_heading("4.2.2 Software Requirements", level=3)
table(["Software", "Version", "Role"],
    [["Windows", "10 or later (any XAMPP-capable OS)", "Host operating system"],
     ["XAMPP", "8.x", "Apache web server, PHP runtime and MySQL bundled together"],
     ["PHP", "8.1 or later", "Server-side language"],
     ["MySQL", "8.0", "Relational database"],
     ["Web browser", "Any modern browser", "Client runtime"],
     ["Visual Studio Code", "Latest", "Development editor"],
     ["phpMyAdmin", "Bundled with XAMPP", "Database administration and schema capture"]],
    cap="Software requirements", widths=[1.7, 2.0, 2.7])

doc.add_heading("4.3 Technologies Used and Their Justification", level=2)
table(["Technology", "Justification"],
    [["XAMPP platform", "Free, widely taught, and installable by a non-technical user. Deployment of the finished system is a folder copy into htdocs; there is no build step and no command line."],
     ["PHP 8", "Mature and ubiquitous on affordable hosting. Provides bcrypt password hashing, session management and PDO in the standard library — the three pillars of this system’s security."],
     ["MySQL 8 (InnoDB)", "Enforced foreign keys and transactions protect commercial data; phpMyAdmin gives the owner a familiar administration surface; prices are stored as integer cents to avoid floating-point error."],
     ["HTML5 / CSS3", "Semantic markup with a hand-written theme keeps the storefront fast and fully under the project’s control; no CSS framework is required."],
     ["Vanilla JavaScript", "Powers the cart, the live bouquet preview, toasts, modal dialogs, confirmation dialogs and form enhancements without any framework or build tooling, preserving the copy-folder deployment property."],
     ["Chart.js (CDN)", "Renders the administrator’s revenue and order-status charts as genuine interactive charts."],
     ["Lucide icons (CDN)", "Provides a consistent, professional icon set across the storefront and dashboard."],
     ["Wikimedia Commons imagery", "Openly licensed flower photography with documented attribution serves as production-quality placeholder imagery until the boutique substitutes its own photographs."]],
    cap="Technologies used and their justification", widths=[1.9, 4.5])

doc.add_heading("4.4 Implementation of the Main Modules", level=2)
doc.add_heading("4.4.1 Automatic Installation and Seeding", level=3)
p("The single most important implementation decision for this case study is that the "
  "system installs itself. On the first request, includes/db.php connects to MySQL, "
  "creates the peonify database if it is missing, creates all twelve tables, and seeds "
  "the administrator account, the catalogue taxonomy, ten demonstration products and "
  "the bouquet-builder options. Every step is idempotent, so the same code path serves "
  "both first installation and every subsequent request without harm.")
doc.add_heading("4.4.2 Storefront and Merchandising", level=3)
p("The landing page implements the merchandising surface designed in Chapter 3: a hero "
  "section with a three-dimensional parallax photograph stack that tilts toward the "
  "pointer; a flash-sale banner linking to discounted arrangements; best sellers "
  "computed from actual order lines; newest arrivals; collection tiles whose cover "
  "images are drawn from the first product of each collection; verified customer "
  "testimonials; and a four-step “how it works” explanation. The shop page composes "
  "search, category, collection, price-range and sale filters with sorting and "
  "pagination, all expressed as ordinary hyperlink and form parameters so the state is "
  "bookmarkable.")
doc.add_heading("4.4.3 Bouquet Builder", level=3)
p("The builder renders its options from the builder_options table in four steps. A "
  "vanilla-JavaScript preview composes circular crops of real flower photography into a "
  "bouquet dome as selections are made, while a running breakdown prices each step. "
  "The completed configuration travels with the cart line as JSON and is stored on the "
  "order item, so the administrator sees exactly what was designed.")
doc.add_heading("4.4.4 Accounts, Profiles and Sessions", level=3)
p("Registration and login use split-screen pages with password visibility toggles and "
  "registration-time password confirmation. Passwords are hashed with bcrypt at cost "
  "twelve; the session identifier is regenerated at login; and the session cookie is "
  "HTTP-only and SameSite. The profile page provides a click-to-upload photograph with "
  "live preview, contact details, and the saved delivery address and city that pre-fill "
  "checkout; password change requires the current password and a matching confirmation.")
doc.add_heading("4.4.5 Cart, Checkout and Payment Recording", level=3)
p("The cart lives in the browser’s localStorage, so it survives navigation and login. "
  "Checkout requires an account and captures recipient, delivery date, time window and "
  "an optional gift note. On submission the server verifies the anti-CSRF token, "
  "recomputes every price from the database — treating the browser’s figures as a "
  "proposal only — creates the order, its items and its first event inside a "
  "transaction, records payment, marks the order paid and issues a unique reference of "
  "the form PNY-XXXXXX. Both parties are notified immediately.")
doc.add_heading("4.4.6 Fulfilment, Notifications and Reminders", level=3)
p("The administrator’s Orders screen provides search, a status filter and pagination; "
  "each paid order carries a single Deliver action guarded by a confirmation dialog. "
  "Symmetrically, the customer’s account offers “I received it” on any paid order. "
  "Either action completes the order, appends an audit event and notifies the other "
  "party. Delivery reminders are generated by a lazy scheduler that runs on page loads "
  "at most once every ten minutes — an approach chosen deliberately because XAMPP "
  "installations have no reliable cron facility — and notify the customer twenty-four "
  "hours and five hours before the chosen window.")
doc.add_heading("4.4.7 Administration Dashboard", level=3)
p("The administrator works in a dedicated workspace with a sidebar of nine sections: "
  "Dashboard, Orders, Products, Catalog, Feedback, Support, Notifications, Activity and "
  "Profile, plus a one-click return to the storefront and sign-out. The dashboard "
  "charts thirty days of revenue and the paid-versus-delivered pipeline with Chart.js "
  "and lists top sellers; Products offers create, edit and delete through a modal form "
  "with a click-to-upload photograph square and a promotional discount field; Catalog "
  "manages categories and collections; Feedback and Support present moderation and the "
  "message inbox; Activity is the audit table of all order events.")

doc.add_heading("4.5 Security Implementation", level=2)
bullets([
    "Authentication: passwords hashed with bcrypt (cost 12); plaintext never stored or "
    "logged; session identifier regenerated at login; sessions carried in HTTP-only, "
    "SameSite cookies inaccessible to page scripts.",
    "Authorisation: every protected page verifies the session and the role; "
    "administrators are excluded from purchasing at both interface and server level; "
    "customers can act only on their own orders, notifications and profile.",
    "Request integrity: an anti-CSRF token is embedded in and verified on every form "
    "submission.",
    "Injection defence: all database access uses PDO prepared statements; no query is "
    "ever assembled from raw user input.",
    "Upload safety: uploads are restricted to image MIME types verified server-side, "
    "limited to five megabytes, and stored under randomised names.",
    "Commercial integrity: order totals are recomputed on the server from the database "
    "at checkout; client-side prices are never trusted.",
])

doc.add_heading("4.6 Presentation of the Developed System", level=2)
p("The figures that follow are authentic captures of the running system, taken from "
  "the deployed application without retouching.")

SHOTS_LIST = [
    ("home", "Landing page — hero with parallax photograph stack, trust points and entry to the shop and builder",
     "The landing page establishes the boutique’s identity in the first screen: the "
     "serif display type, the rose-and-ivory palette and authentic photography, with "
     "the two principal calls to action and the floating customer testimonial."),
    ("shop", "Shop page — category chips, filters, search, sorting and the paginated product grid",
     "All discovery criteria compose freely and are expressed in the address bar, so "
     "any filtered view can be bookmarked or shared. Discount and “New” badges are "
     "computed from the catalogue data."),
    ("product", "Product page — discounted pricing, add-to-cart and verified customer feedback",
     "The product page presents the effective price alongside the original when a "
     "promotion applies, and hosts the feedback section where verified customers rate "
     "and review the arrangement."),
    ("builder", "Bouquet builder — four steps with a live photographic preview and price breakdown",
     "Each selection updates the composed preview and the per-step breakdown; the "
     "completed design is added to the cart with its configuration attached."),
    ("cart", "Cart — quantity controls, custom-bouquet lines and the checkout call to action",
     "The cart persists in the browser and distinguishes catalogue products from "
     "custom bouquets; totals update instantly as quantities change."),
    ("login", "Login page — split-screen design with password visibility toggle",
     "Authentication pages pair the form with full-bleed photography; field icons and "
     "the visibility toggle reduce input errors."),
    ("register", "Registration page — with password confirmation",
     "Registration validates email format, minimum password length and the match "
     "between password and confirmation before an account is created."),
    ("checkout", "Checkout — recipient details, delivery scheduling and the order summary",
     "Recipient details arrive pre-filled from the profile; the customer chooses "
     "on-demand or a future date with a time window; the summary presents the exact "
     "server-verified total."),
    ("account", "Customer dashboard — status cards and paginated orders with status badges",
     "The customer sees at a glance how many orders are awaiting delivery and "
     "delivered, the total spent, and each order’s state, with details and delivery "
     "confirmation one click away."),
    ("account_profile", "Customer profile — click-to-upload photograph, saved delivery details and password change",
     "The saved address and city pre-fill every future checkout; password change "
     "requires the current password and a matching confirmation."),
    ("contact", "Contact page — the public form that feeds the administrator’s inbox",
     "Messages submitted here create administrator notifications and appear in the "
     "Support section of the dashboard."),
    ("admin_dashboard", "Administrator dashboard — headline statistics, thirty-day revenue chart and order pipeline",
     "The dashboard answers the owner’s daily questions in one screen: revenue "
     "movement, orders awaiting delivery, catalogue size and customer count, with top "
     "sellers ranked by units sold."),
    ("admin_orders", "Administrator orders — search, status filter and the single Deliver action",
     "Every order arrives already paid; the administrator’s only task is the "
     "confirmed Deliver action, after which the customer is notified automatically."),
    ("admin_products", "Administrator products — the modal product form with click-to-upload photograph",
     "Products are created and edited in a modal dialog; the photograph square "
     "previews the upload before saving, and the discount field powers the "
     "storefront’s flash-sale merchandising."),
    ("admin_catalog", "Administrator catalog — categories and collections management",
     "The taxonomy the customers filter by is fully under the owner’s control; new "
     "entries appear in the storefront immediately."),
    ("admin_activity", "Administrator activity — the audit table of all order events",
     "Every order transition, with its timestamp and note, is permanently visible — "
     "the accountability record the manual process never had."),
]
for name, cap, desc in SHOTS_LIST:
    screenshot(name, cap)
    p(desc)

doc.add_heading("4.7 System Testing", level=2)
doc.add_heading("4.7.1 Testing Approach", level=3)
p("Testing proceeded at three complementary levels throughout development. Unit-level "
  "checks validated individual behaviours — price computation with discounts, reference "
  "generation, reminder timing arithmetic. Integration testing exercised complete "
  "request cycles against a live MySQL instance, verifying that each form submission "
  "produced the correct database state, notifications and redirects. Finally, "
  "system-level acceptance testing walked every user journey end to end in the "
  "browser, on both desktop and mobile viewports, including the failure paths: wrong "
  "passwords, mismatched confirmations, tampered tokens, oversized uploads and "
  "attempts to act on another user’s data.")
doc.add_heading("4.7.2 Test Cases and Results", level=3)
table(["#", "Test case", "Expected result", "Result"],
    [["1", "Register with mismatched passwords", "Error message; account not created", "Pass"],
     ["2", "Register and sign in", "Session starts; dashboard greets the user by name", "Pass"],
     ["3", "Sign in with a wrong password", "Error message; no session created", "Pass"],
     ["4", "Guest attempts checkout", "Redirected to login; returned to checkout after signing in", "Pass"],
     ["5", "Administrator attempts to buy", "Add-to-cart refused; checkout blocks administrator accounts", "Pass"],
     ["6", "Checkout with catalogue and custom items", "Order created as paid; totals recomputed server-side; unique reference issued", "Pass"],
     ["7", "Cart price manipulation attempt", "Server total derived from the database, ignoring altered client prices", "Pass"],
     ["8", "Administrator delivers an order", "Confirmation dialog; status delivered; customer notified", "Pass"],
     ["9", "Customer confirms receipt", "Status delivered; administrator notified", "Pass"],
     ["10", "Customer acts on another customer’s order", "Action refused", "Pass"],
     ["11", "Delivery reminders", "Notifications generated 24 hours and 5 hours before the window", "Pass"],
     ["12", "Post, update and moderate feedback", "One review per customer per product; appears on product and landing pages; administrator can remove", "Pass"],
     ["13", "Contact form submission", "Message stored in the inbox; administrator notified", "Pass"],
     ["14", "CSRF token tampering", "Request rejected", "Pass"],
     ["15", "Oversized / non-image upload", "Upload rejected with an error message", "Pass"],
     ["16", "First-run installation", "Database, tables, administrator account and demonstration catalogue created automatically", "Pass"]],
    cap="Functional test cases and results", widths=[0.4, 2.1, 3.2, 0.6])
pagebreak()

# ╔══════════════════════════ CHAPTER 5 ══════════════════════════╗
doc.add_heading("CHAPTER 5", level=1)
doc.add_heading("CONCLUSION AND RECOMMENDATIONS", level=1)

doc.add_heading("5.1 Summary of the Work Done", level=2)
p("This project set out to replace the fragile, chat-based selling process of a small "
  "floral boutique with complete, professional commerce infrastructure. The existing "
  "manual process was studied and modelled; its problems were traced to the absence of "
  "durable records at every stage; and a system was designed with object-oriented "
  "techniques, implemented on the PHP/MySQL/XAMPP stack, and tested across every user "
  "journey. The delivered system comprises a merchandised storefront, an interactive "
  "bouquet builder, secure role-based accounts, a server-verified checkout with "
  "payment recording and unique order references, a single-action fulfilment workflow "
  "with notifications and delivery reminders, verified product feedback, and an "
  "administrator dashboard with genuine analytics.")

doc.add_heading("5.2 Achievement of Objectives", level=2)
table(["Specific objective", "Achievement"],
    [["Analyse and model the existing manual process", "Achieved — Chapter 2 presents the environment, the process model (Figure 1) and the derived problems."],
     ["Design a premium storefront with a bouquet builder", "Achieved — implemented with live photographic preview and full merchandising (Figures 9–12)."],
     ["Implement secure role-based accounts", "Achieved — bcrypt, regenerated sessions, HTTP-only cookies, CSRF protection, separated dashboards."],
     ["Implement checkout with server-side pricing and unique references", "Achieved — transactional order creation; manipulation attempts defeated in testing (Test 7)."],
     ["Implement the fulfilment workflow with notifications and reminders", "Achieved — single-action delivery by either party; reminders at 24 h and 5 h (Tests 8–11)."],
     ["Provide analytics and catalogue control", "Achieved — Chart.js dashboard, product CRUD with photographs and discounts, catalogue management (Figures 20–23)."],
     ["Test every user journey", "Achieved — sixteen documented test cases, all passing (Table 24)."]],
    cap="Achievement of the specific objectives", widths=[3.2, 3.2])

doc.add_heading("5.3 Challenges Encountered", level=2)
p("Three challenges shaped the final system. First, designing for a non-technical "
  "operator required repeatedly simplifying features that were technically satisfying "
  "but operationally heavy — most visibly the reduction of the fulfilment pipeline to "
  "a single transition. Second, the absence of a scheduling facility on XAMPP required "
  "an alternative design for delivery reminders, solved with a lazy scheduler paced "
  "through the database. Third, guaranteeing commercial integrity against a "
  "manipulable browser required moving every price computation to the server and "
  "treating the client cart strictly as a proposal.")

doc.add_heading("5.4 Recommendations", level=2)
bullets([
    "The boutique should replace the openly licensed placeholder photography with its "
    "own product photographs before commercial launch.",
    "The default administrator credentials and the application’s secret values should "
    "be changed immediately after installation.",
    "A TLS certificate should be configured when the system moves from local testing "
    "to public hosting.",
    "Regular database backups should be scheduled using phpMyAdmin’s export facility "
    "or mysqldump.",
])

doc.add_heading("5.5 Future Work", level=2)
bullets([
    "Integration of an online payment gateway (card and mobile money) at checkout.",
    "Delivery of the existing notifications additionally by email and SMS.",
    "Order cancellation and refund handling with the corresponding audit events.",
    "Inventory quantity tracking per product with low-stock alerts.",
    "A multi-shop extension of the same role model, allowing several boutiques to "
    "share one installation.",
])
pagebreak()

# ╔══════════════════════════ REFERENCES ══════════════════════════╗
doc.add_heading("REFERENCES", level=1)
p("All online sources were accessed in July 2026.", italic=True)
REFS = [
    "1-800-Flowers.com, Inc. (2026). Flowers, flower delivery, fresh flowers online. https://www.1800flowers.com/",
    "Apache Friends. (2026). XAMPP Apache + MariaDB + PHP + Perl. https://www.apachefriends.org/",
    "Bloom & Wild. (2026). Flower delivery — letterbox flowers and gifts. https://www.bloomandwild.com/",
    "BusinessDojo. (2025). Flower shop industry: Market statistics and trends. https://dojobusiness.com/blogs/news/flower-shop-industry-statistics",
    "Chart.js. (2026). Chart.js — simple yet flexible JavaScript charting (Documentation). https://www.chartjs.org/docs/",
    "Floward. (2026). Online flowers & gifts — same-day flower delivery. https://floward.com/",
    "Grand View Research. (2024). Flower delivery service market size & share report, 2025–2030. https://www.grandviewresearch.com/industry-analysis/flower-delivery-service-market-report",
    "Lucide. (2026). Lucide — beautiful & consistent icons (Documentation). https://lucide.dev/",
    "Mastercard Strive. (2024). Social commerce: How micro and small businesses sell on social platforms. https://strivecommunity.org/insights/market-participation/social-commerce",
    "MDN Web Docs. (2026). Using HTTP cookies — HttpOnly, SameSite and Secure attributes. Mozilla. https://developer.mozilla.org/en-US/docs/Web/HTTP/Cookies",
    "MySQL. (2026). MySQL 8.0 reference manual. Oracle Corporation. https://dev.mysql.com/doc/refman/8.0/en/",
    "OWASP Foundation. (2025). Cross-site request forgery prevention cheat sheet. https://cheatsheetseries.owasp.org/cheatsheets/Cross-Site_Request_Forgery_Prevention_Cheat_Sheet.html",
    "OWASP Foundation. (2025). Password storage cheat sheet. https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html",
    "PHP Group. (2026). PHP 8 manual — PDO, sessions and password hashing. https://www.php.net/manual/en/",
    "PwC. (2024). Is your brand ready for the $3 trillion social commerce marketplace? https://www.pwc.com/us/en/services/consulting/business-transformation/library/social-commerce.html",
    "Sommerville, I. (2016). Software Engineering (10th ed.). Pearson Education.",
    "Tidio. (2024). 11 key social commerce statistics to know. https://www.tidio.com/blog/social-commerce-statistics/",
    "Wikimedia Foundation. (2026). Wikimedia Commons — reusing content outside Wikimedia. https://commons.wikimedia.org/wiki/Commons:Reusing_content_outside_Wikimedia",
]
for r in sorted(REFS, key=str.lower):
    par = doc.add_paragraph(r)
    par.paragraph_format.left_indent = Inches(0.4)
    par.paragraph_format.first_line_indent = Inches(-0.4)
    par.paragraph_format.space_after = Pt(6)
pagebreak()

# ╔══════════════════════════ APPENDICES ══════════════════════════╗
doc.add_heading("APPENDICES", level=1)
doc.add_heading("Appendix A: Installation Guide (Windows / XAMPP)", level=2)
numbered([
    "Install XAMPP from apachefriends.org and open the XAMPP Control Panel.",
    "Press Start next to Apache and next to MySQL.",
    "Copy the peonify-php folder into C:\\xampp\\htdocs\\ and rename it to peonify.",
    "Open http://localhost/peonify/ in a browser. On this first visit the system "
    "creates the database, all tables, the administrator account and the "
    "demonstration catalogue automatically.",
    "Sign in as the administrator at http://localhost/peonify/admin/ with "
    "admin@peonify.com / peonify-admin, and change the password from the Profile "
    "section immediately.",
])
doc.add_heading("Appendix B: Configuration Reference (config.php)", level=2)
table(["Constant", "Purpose", "Default"],
    [["DB_HOST / DB_PORT", "MySQL server location", "127.0.0.1 / 3306"],
     ["DB_NAME", "Database name (created automatically)", "peonify"],
     ["DB_USER / DB_PASS", "MySQL credentials (XAMPP defaults)", "root / (empty)"],
     ["ADMIN_EMAIL / ADMIN_PASSWORD", "Seeded administrator account", "admin@peonify.com / peonify-admin"]],
    cap="Configuration reference", widths=[2.0, 2.6, 1.8])
doc.add_heading("Appendix C: Representative Code Excerpts", level=2)
p("Password verification and session establishment (login.php):", italic=True)
p("if ($user && password_verify($_POST['password'], $user['password_hash'])) {\n"
  "    session_regenerate_id(true);\n"
  "    $_SESSION['uid'] = (int)$user['id'];\n"
  "}", align="left", size=10)
p("Server-side price recomputation at checkout (checkout.php):", italic=True)
p("$st = $pdo->prepare('SELECT * FROM products WHERE id = ? AND in_stock = 1');\n"
  "$st->execute([(int)$item['product_id']]);\n"
  "$unit = effective_price($st->fetch());   // client price is ignored",
  align="left", size=10)
p("Anti-CSRF verification applied to every form (functions.php):", italic=True)
p("function csrf_check(): void {\n"
  "    if (($_POST['csrf'] ?? '') !== ($_SESSION['csrf'] ?? null)) {\n"
  "        http_response_code(400); exit('Invalid request token.');\n"
  "    }\n"
  "}", align="left", size=10)
doc.add_heading("Appendix D: Image Attribution", level=2)
p("All storefront and builder photography is sourced from Wikimedia Commons under free "
  "licences; the complete per-file attribution list ships with the system at "
  "peonify-php/assets/images/CREDITS.md. The images are placeholders to be replaced "
  "with the boutique’s own photography before commercial launch.")

# ---- fill the lists of figures and tables -----------------------------------
LOF_PAR.text = ""
for f in FIGS:
    LOF_PAR.insert_paragraph_before(f, style="List Bullet")
LOT_PAR.text = ""
for t in TABS:
    LOT_PAR.insert_paragraph_before(t, style="List Bullet")

doc.save(OUT)
print("Saved:", OUT)
print("figures:", len(FIGS), "| tables:", len(TABS))
